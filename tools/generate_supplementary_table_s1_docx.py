"""
Generate only Supplementary Table S1 (Hallmark ORA) as a standalone Word file.

Publication-oriented layout: core columns only (no per-row gene lists in the grid—those
stay in the source CSV), explicit column widths, compact font, and header row styling.

Usage (from feasibility_study/):
  python tools/generate_supplementary_table_s1_docx.py
  python tools/generate_supplementary_table_s1_docx.py --outputs-dir E:\\path\\to\\outputs

Output:
  outputs/manuscript/Supplementary_Table_S1.docx
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

import pandas as pd

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError as e:
    print("Install python-docx: pip install python-docx", file=sys.stderr)
    raise SystemExit(1) from e


PROJECT_ROOT = Path(__file__).resolve().parents[1]


def _safe_csv(p: Path) -> pd.DataFrame | None:
    if not p.is_file():
        return None
    try:
        return pd.read_csv(p, nrows=None)
    except Exception:
        return None


def _fmt_p(x: object) -> str:
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return ""
    try:
        v = float(x)
    except (TypeError, ValueError):
        return str(x)[:32]
    if v == 0.0:
        return "0"
    if abs(v) < 1e-6 or abs(v) >= 1e4:
        return f"{v:.3e}"
    return f"{v:.6g}".rstrip("0").rstrip(".")


def _fmt_float(x: object, nd: int = 4) -> str:
    if x is None or (isinstance(x, float) and pd.isna(x)):
        return ""
    try:
        v = float(x)
    except (TypeError, ValueError):
        return str(x)[:24]
    if abs(v) >= 1e4 or (abs(v) > 0 and abs(v) < 1e-3):
        return f"{v:.{nd}e}"
    return f"{v:.{nd}f}".rstrip("0").rstrip(".")


def _set_cell_text(cell, text: str, *, bold: bool = False, size_pt: float = 9, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"


def _set_cell_width(cell, width_twips: int) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for el in tcPr.findall(qn("w:tcW")):
        tcPr.remove(el)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_table_grid(table, col_widths_twips: list[int]) -> None:
    """Define column widths. tblGrid MUST follow tblPr before any w:tr (ECMA-376); Word rejects if tblGrid is last."""
    tbl = table._tbl
    for el in list(tbl.findall(qn("w:tblGrid"))):
        tbl.remove(el)
    tbl_grid = OxmlElement("w:tblGrid")
    for w in col_widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(w))
        tbl_grid.append(grid_col)
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl.insert(0, tbl_grid)
        return
    insert_at = list(tbl).index(tbl_pr) + 1
    tbl.insert(insert_at, tbl_grid)


def _set_table_full_width(table, width_twips: int) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    for el in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(el)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    for el in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(el)
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)


def _repeat_header_row(table) -> None:
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    for el in trPr.findall(qn("w:tblHeader")):
        trPr.remove(el)
    tbl_header = OxmlElement("w:tblHeader")
    trPr.append(tbl_header)


def _shade_header_row(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9E2F3")
        shading.set(qn("w:val"), "clear")
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_pr.append(shading)


def build_table_s1_docx(out_dir: Path, doc_path: Path, *, n_terms: int = 15) -> None:
    out_dir = out_dir.resolve()
    csv_path = out_dir / "enrichr_hallmark_ora_mirtarbase_union.csv"
    df = _safe_csv(csv_path)
    if df is None or df.empty:
        print(f"[error] Missing or empty: {csv_path}", file=sys.stderr)
        raise SystemExit(1)

    need = {"Term", "Overlap", "P-value", "Adjusted P-value", "Odds Ratio", "Combined Score"}
    if not need.issubset(set(df.columns)):
        print(f"[error] CSV missing expected columns. Have: {list(df.columns)}", file=sys.stderr)
        raise SystemExit(1)

    view = df.head(n_terms).copy()
    rows_out: list[tuple[str, str, str, str, str, str, str]] = []
    for i, (_, r) in enumerate(view.iterrows(), start=1):
        rows_out.append(
            (
                str(i),
                str(r.get("Term", ""))[:120],
                str(r.get("Overlap", "")),
                _fmt_p(r.get("P-value")),
                _fmt_p(r.get("Adjusted P-value")),
                _fmt_float(r.get("Odds Ratio")),
                _fmt_float(r.get("Combined Score"), nd=3),
            )
        )

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "Supplementary Table S1 | MSigDB Hallmark over-representation (ORA) on miRTarBase target union",
        level=1,
    )
    title.runs[0].font.size = Pt(14)

    doc.add_paragraph()
    p_src = doc.add_paragraph()
    r = p_src.add_run(
        f"Source: {csv_path.name} (repository outputs). "
        f"Showing top {len(rows_out)} terms by Enrichr / MSigDB Hallmark 2020 ORA. "
        "Gene membership lists are not duplicated here to preserve row height; see the Genes column in the source CSV."
    )
    r.font.size = Pt(9)
    r.italic = True

    doc.add_paragraph()

    headers = (
        "Rank",
        "Hallmark term",
        "Overlap (k/N)",
        "P value",
        "Adjusted P",
        "Odds ratio",
        "Combined score",
    )
    # Total usable text width ~ landscape letter minus margins: ~10 in -> 14400 twips
    col_w = [520, 4680, 1100, 1180, 1180, 1020, 1180]

    table = doc.add_table(rows=1 + len(rows_out), cols=len(headers))
    table.style = "Table Grid"
    _set_table_full_width(table, sum(col_w))
    _set_table_grid(table, col_w)
    _repeat_header_row(table)

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        _set_cell_width(cell, col_w[j])
        _set_cell_text(cell, h, bold=True, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row_data in enumerate(rows_out, start=1):
        for j, val in enumerate(row_data):
            cell = table.rows[ri].cells[j]
            _set_cell_width(cell, col_w[j])
            align = WD_ALIGN_PARAGRAPH.RIGHT if j >= 2 else (WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_text(cell, val, size_pt=9, align=align)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tr = table.rows[ri]._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), "400")
            trHeight.set(qn("w:hRule"), "atLeast")
            for old in trPr.findall(qn("w:trHeight")):
                trPr.remove(old)
            trPr.append(trHeight)

    _shade_header_row(table)

    doc.add_paragraph()
    leg = doc.add_paragraph()
    lr = leg.add_run("Table legend. ")
    lr.bold = True
    lr.font.size = Pt(10)
    leg.add_run(
        "Overlap lists k genes in the miRTarBase-derived target query that appear in each Hallmark set of size N. "
        "P values are from hypergeometric over-representation; Adjusted P is Benjamini–Hochberg FDR within the ORA output table. "
        "Odds ratio and combined score follow Enrichr conventions. Full columns (including Genes) appear in enrichr_hallmark_ora_mirtarbase_union.csv."
    ).font.size = Pt(9)

    doc_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(doc_path))
    print(f"Wrote {doc_path}")


def main() -> int:
    ap = argparse.ArgumentParser(description="Standalone Supplementary Table S1 (.docx) only.")
    ap.add_argument(
        "--outputs-dir",
        type=Path,
        default=PROJECT_ROOT / "outputs",
        help="Directory containing enrichr_hallmark_ora_mirtarbase_union.csv",
    )
    ap.add_argument(
        "--out-doc",
        type=Path,
        default=None,
        help="Destination .docx (default: <outputs-dir>/manuscript/Supplementary_Table_S1.docx)",
    )
    ap.add_argument("--n-terms", type=int, default=15, help="Number of top rows to include (default: 15).")
    args = ap.parse_args()
    out_dir = args.outputs_dir.resolve()
    doc_path = args.out_doc or (out_dir / "manuscript" / "Supplementary_Table_S1.docx")
    build_table_s1_docx(out_dir, doc_path, n_terms=args.n_terms)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
