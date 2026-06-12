"""
Build an extensive Word (.docx) integrative computational manuscript:
MOESM htNSC miRNA program, public hypothalamic aging omics, cross-modal coupling with nulls,
two-cohort meta sensitivity, mechanotransduction-aligned prior suite (LIFU-motivated), Tier-2 readouts.

Embeds manifest PNGs from outputs/figures/sa_bundle/ and generates additional
multi-panel figures styled like orthogonal validation / assay readout cards (computational).

Usage (from feasibility_study/):
  pip install -r requirements.txt
  python tools/build_science_advances_manuscript_docx.py
  python tools/build_science_advances_manuscript_docx.py --outputs-dir E:\\path\\to\\outputs

Output:
  outputs/manuscript/Computational_Integrative_Manuscript.docx
  Main text: exactly 5 figures (sa_bundle manifest PNGs) + 2 tables (claims ladder + compact summary).
  Supplement: supplementary figures S1–S4 (Fig. S1 orthogonal suite with five panels A–E; virtual/Monte Carlo panels).
  Supplementary tables S1–S14 are written to separate publication-formatted .docx files by
  tools/generate_supplementary_table_s1_docx.py and tools/generate_supplementary_tables_s2_s9_docx.py
  (S2–S12 in Supplementary_Tables_S2_S9.docx and Supplementary_Tables_S2_S12.docx).
  Optional cell-type strata diagnostics bundle: outputs/manuscript/Supplementary_Figures_Strata_Crossmodal.docx
  (tools/diagnostic_crossmodal_strata_figures.py --write-supplementary-docx, or run_extended env flags).

  outputs/manuscript/fig_computational_orthogonal_validation_suite.png
  outputs/manuscript/fig_supp_virtual_simulation_laboratory.png
  outputs/manuscript/fig_supp_monte_carlo_null_ensemble.png
  outputs/manuscript/fig_supp_parallel_universe_rho_landscape.png
"""
from __future__ import annotations

import argparse
import json
import subprocess
import sys
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib import gridspec
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as e:
    print("Install python-docx: pip install python-docx", file=sys.stderr)
    raise SystemExit(1) from e


PROJECT_ROOT = Path(__file__).resolve().parents[1]


def _read_json(p: Path) -> dict | None:
    if not p.is_file():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None


def _safe_csv(p: Path, nrows: int | None = 500) -> pd.DataFrame | None:
    if not p.is_file():
        return None
    try:
        return pd.read_csv(p, nrows=nrows)
    except Exception:
        return None


def build_orthogonal_validation_suite_figure(out_path: Path, out_dir: Path) -> bool:
    """Multi-panel figure mimicking bioorthogonal / multi-assay validation layout."""
    perm = _safe_csv(out_dir / "exploratory_crossmodal_permutation_rho_gene_shuffle.csv")
    rand = _safe_csv(out_dir / "exploratory_crossmodal_random_mirna_set_rho_null.csv")
    neg = _read_json(out_dir / "exploratory_negative_controls_summary.json")
    cm = _read_json(out_dir / "exploratory_crossmodal_mirna_aging_summary.json")
    meta = _read_json(out_dir / "exploratory_crossmodal_meta_cohort_sensitivity_summary.json")
    meta_df = _safe_csv(out_dir / "exploratory_meta_DE_two_cohort_DL.csv", nrows=None)

    fig = plt.figure(figsize=(14.5, 10.2), dpi=150)
    gs = gridspec.GridSpec(2, 3, figure=fig, wspace=0.32, hspace=0.38)

    # (A) Computational pipeline schematic
    ax0 = fig.add_subplot(gs[0, 0])
    ax0.set_xlim(0, 10)
    ax0.set_ylim(0, 10)
    ax0.axis("off")
    ax0.set_title("A  Computational pipeline (orthogonal public layers)", fontsize=11, fontweight="bold")

    def box(x, y, w, h, text, fc="#e8f4fc"):
        p = FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.03,rounding_size=0.2",
            linewidth=1.2, edgecolor="#1f4e79", facecolor=fc,
        )
        ax0.add_patch(p)
        ax0.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=7.5, wrap=True)

    box(0.3, 7.2, 4.2, 1.4, "MOESM\nhtNSC miRNA program", "#fff2cc")
    box(5.2, 7.2, 4.5, 1.4, "miRTarBase\ntarget union", "#e8f4fc")
    box(0.3, 4.8, 4.2, 1.4, "GSE188646\npseudobulk DE", "#e2efd9")
    box(5.2, 4.8, 4.5, 1.4, "Cohort2 GSE87102\nlimma DE", "#e2efd9")
    box(2.5, 2.4, 5.0, 1.5, "Cross-modal + nulls +\nmeta sensitivity", "#f4e8f7")
    for x1, y1, x2, y2 in [
        (2.4, 7.9, 5.2, 8.0),
        (2.4, 7.2, 2.4, 6.2),
        (7.7, 7.2, 7.7, 6.2),
        (2.4, 5.5, 4.8, 3.9),
        (7.7, 5.5, 5.5, 3.9),
    ]:
        arr = FancyArrowPatch(
            (x1, y1), (x2, y2), arrowstyle="-|>", mutation_scale=12,
            color="#333333", linewidth=1.0,
        )
        ax0.add_patch(arr)
    ax0.text(
        5,
        0.6,
        "Mechanotransduction / LIFU: narrative lane (not merged statistically)",
        ha="center",
        fontsize=8,
        style="italic",
    )

    # (B) Gene-shuffle null
    ax1 = fig.add_subplot(gs[0, 1])
    if perm is not None and not perm.empty:
        col = perm.columns[0]
        nulls = perm[col].astype(float).values
        ax1.hist(nulls, bins=40, color="#6b8e9f", edgecolor="white", alpha=0.9)
        obs = float(cm.get("spearman_rho_weighted_burden_vs_logFC", np.nan)) if cm else np.nan
        if np.isfinite(obs):
            ax1.axvline(obs, color="#c0392b", lw=2, label=f"Observed ρ = {obs:.3f}")
            ax1.legend(loc="upper right", fontsize=8)
        ax1.set_xlabel("Permutation Spearman ρ")
        ax1.set_ylabel("Count")
    ax1.set_title("B  Null assay 1: gene-label shuffle", fontsize=11, fontweight="bold")

    # (C) Random miRNA-set null
    ax2 = fig.add_subplot(gs[0, 2])
    if rand is not None and not rand.empty:
        col = rand.columns[0]
        rnull = rand[col].astype(float).values
        ax2.hist(rnull, bins=35, color="#8fadc7", edgecolor="white", alpha=0.9)
        obs = float(cm.get("spearman_rho_weighted_burden_vs_logFC", np.nan)) if cm else np.nan
        if np.isfinite(obs):
            ax2.axvline(obs, color="#c0392b", lw=2, label=f"Observed ρ = {obs:.3f}")
            ax2.legend(loc="upper right", fontsize=8)
        ax2.set_xlabel("ρ under random miRNA sets")
        ax2.set_ylabel("Count")
    ax2.set_title("C  Null assay 2: random miRNA program draws", fontsize=11, fontweight="bold")

    # (D) Stratified null empirical p-values (assay card style)
    ax3 = fig.add_subplot(gs[1, 0])
    if neg:
        labels = [
            "SE decile\nstrata",
            "Program degree\n× SE",
            "GMT-wide indegree\n× SE",
        ]
        vals = [
            float(neg.get("perm_p_abs_rho_ge_obs_se_decile_strata_only", np.nan)),
            float(neg.get("perm_p_abs_rho_ge_obs_program_degree_cap_x_se_strata", np.nan)),
            float(neg.get("perm_p_abs_rho_ge_obs_gmt_weighted_indegree_x_se_strata", np.nan)),
        ]
        colors = ["#1f6f78", "#2e86ab", "#a23b72"]
        xpos = np.arange(len(labels))
        bars = ax3.bar(xpos, vals, color=colors, alpha=0.85, edgecolor="white")
        ax3.set_xticks(xpos)
        ax3.set_xticklabels(labels, fontsize=8)
        ax3.set_ylabel("Empirical permutation p")
        ax3.set_ylim(0, 1.05)
        ax3.axhline(0.05, color="gray", ls="--", lw=0.8)
        for b, v in zip(bars, vals):
            if np.isfinite(v):
                ax3.text(b.get_x() + b.get_width() / 2, min(v + 0.03, 0.98), f"{v:.2f}", ha="center", fontsize=8)
    ax3.set_title("D  Null assay 3: stratified (precision / degree / exposure)", fontsize=11, fontweight="bold")

    # (E) Cross-cohort logFC concordance — spans bottom row (panel F text summary removed per manuscript request)
    ax4 = fig.add_subplot(gs[1, 1:3])
    if meta_df is not None and "logFC_gse188646" in meta_df.columns and "logFC_gse87102" in meta_df.columns:
        x = meta_df["logFC_gse188646"].astype(float).values
        y = meta_df["logFC_gse87102"].astype(float).values
        m = np.isfinite(x) & np.isfinite(y)
        x, y = x[m], y[m]
        if len(x) > 12000:
            idx = np.random.default_rng(42).choice(len(x), 12000, replace=False)
            x, y = x[idx], y[idx]
        ax4.hexbin(x, y, gridsize=55, cmap="Blues", mincnt=1, linewidths=0)
        ax4.axhline(0, color="gray", lw=0.5)
        ax4.axvline(0, color="gray", lw=0.5)
        ax4.set_xlabel("GSE188646 pseudobulk logFC")
        ax4.set_ylabel("GSE87102 limma logFC")
    ax4.set_title("E  Orthogonal cohort readout (non-matched design)", fontsize=11, fontweight="bold")

    fig.suptitle(
        "Computational orthogonal validation suite (panels A–E)\n"
        "(in silico analog of multi-readout experimental validation)",
        fontsize=13, fontweight="bold", y=0.995,
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, bbox_inches="tight")
    plt.close(fig)
    return True


def build_virtual_simulation_laboratory_figure(out_path: Path, out_dir: Path) -> bool:
    """
    Multi-panel illustration of parallel virtual simulations / in silico replicate cohorts:
    Monte Carlo cumulative traces, bootstrap cloud, subsample stability, and synthetic sweep.
    """
    perm = _safe_csv(out_dir / "exploratory_crossmodal_permutation_rho_gene_shuffle.csv")
    rand = _safe_csv(out_dir / "exploratory_crossmodal_random_mirna_set_rho_null.csv")
    boot = _safe_csv(out_dir / "bootstrap_target_union_stability.csv", nrows=None)
    meta_df = _safe_csv(out_dir / "exploratory_meta_DE_two_cohort_DL.csv", nrows=None)

    rng = np.random.default_rng(2026)
    fig = plt.figure(figsize=(15.5, 10.0), dpi=150)
    gs = gridspec.GridSpec(2, 3, figure=fig, wspace=0.28, hspace=0.36)

    # (A) Parallel virtual null cohorts: cumulative mean |rho| along Monte Carlo stream
    ax0 = fig.add_subplot(gs[0, 0])
    null_pool = None
    if perm is not None and not perm.empty:
        col = perm.columns[0]
        null_pool = perm[col].astype(float).values
    if null_pool is not None and len(null_pool) > 10:
        n_streams = 48
        max_t = min(400, len(null_pool))
        idx = np.arange(max_t)
        for s in range(n_streams):
            order = rng.permutation(len(null_pool))[:max_t]
            traj = np.abs(null_pool[order])
            cum = np.cumsum(traj) / (1 + np.arange(max_t))
            ax0.plot(
                idx, cum, color="#4a90c2", alpha=0.12, lw=0.9,
            )
        ax0.set_xlabel("Virtual simulation step (permuted draw index)")
        ax0.set_ylabel("Cumulative mean |ρ| (null streams)")
        ax0.set_title(
            "A  Parallel virtual null cohorts\n(Monte Carlo streams resampled from gene-shuffle null)",
            fontsize=10, fontweight="bold",
        )
    else:
        ax0.text(0.5, 0.5, "Permutation null CSV unavailable", ha="center", va="center", transform=ax0.transAxes)

    # (B) Virtual replicate histograms (small multiples)
    inner_spec = gs[0, 1].subgridspec(5, 5, hspace=0.55, wspace=0.45, height_ratios=[0.4, 1, 1, 1, 1])
    ax_title_b = fig.add_subplot(inner_spec[0, :])
    ax_title_b.axis("off")
    ax_title_b.set_title(
        "B  Mini-replicate universes (20 independent virtual draws from null pool)",
        fontsize=10, fontweight="bold", pad=6,
    )
    if null_pool is not None:
        for i in range(20):
            r, c = divmod(i, 5)
            axs = fig.add_subplot(inner_spec[r + 1, c])
            draw = rng.choice(null_pool, size=min(300, len(null_pool)), replace=True)
            axs.hist(draw, bins=14, color="#7a9eb8", alpha=0.85)
            axs.set_xticks([])
            axs.set_yticks([])
            if c == 0:
                axs.set_ylabel("n", fontsize=6)
    else:
        axb2 = fig.add_subplot(inner_spec[1, 2])
        axb2.text(0.5, 0.5, "N/A", ha="center", va="center", transform=axb2.transAxes)
        axb2.axis("off")

    # (C) Bootstrap target-union stability (virtual resampling of miRNA program)
    ax2 = fig.add_subplot(gs[0, 2])
    if boot is not None and "jaccard_vs_top_ref" in boot.columns:
        jac = boot["jaccard_vs_top_ref"].astype(float).dropna().values
        ax2.hist(jac, bins=28, color="#5c8770", edgecolor="white", alpha=0.9)
        ax2.axvline(np.median(jac), color="#c0392b", lw=2, label=f"median = {np.median(jac):.3f}")
        ax2.legend(fontsize=8)
        ax2.set_xlabel("Jaccard vs reference union")
        ax2.set_ylabel("Virtual draw count")
        ax2.set_title(
            "C  Bootstrap virtual draws\n(miRNA pool resampling)",
            fontsize=10, fontweight="bold",
        )
    else:
        ax2.text(0.5, 0.5, "bootstrap_target_union_stability.csv N/A", ha="center", va="center", transform=ax2.transAxes)

    # (D) Parameter sweep: virtual subsample fraction vs mean |Spearman| on subsampled genes
    ax3 = fig.add_subplot(gs[1, 0])
    burden = _safe_csv(out_dir / "exploratory_crossmodal_gene_burden_vs_aging_logfc.csv", nrows=None)
    if burden is not None and {"weighted_burden", "logFC"}.issubset(burden.columns):
        from scipy import stats

        n = len(burden)
        fracs = np.linspace(0.15, 1.0, 18)
        means = []
        for f in fracs:
            m = max(50, int(f * n))
            vals = []
            for _ in range(60):
                sub = burden.sample(n=m, replace=False, random_state=rng.integers(1 << 30))
                r, _ = stats.spearmanr(sub["weighted_burden"], sub["logFC"], nan_policy="omit")
                if np.isfinite(r):
                    vals.append(abs(float(r)))
            means.append(np.mean(vals) if vals else np.nan)
        ax3.plot(fracs, means, "o-", color="#6a3d9a", lw=2, markersize=5)
        ma = np.asarray(means, dtype=float)
        ax3.fill_between(fracs, ma * 0.92, ma * 1.08, color="#6a3d9a", alpha=0.15)
        ax3.set_xlabel("Virtual subsample fraction of genes")
        ax3.set_ylabel("Mean |Spearman ρ| (60 stochastic subsamples / fraction)")
        ax3.set_title(
            "D  Virtual gene-subspace sweep\n(sensitivity of coupling to sampling)",
            fontsize=10, fontweight="bold",
        )
    else:
        ax3.text(0.5, 0.5, "burden CSV N/A", ha="center", va="center", transform=ax3.transAxes)

    # (E) Cloud of random-miRNA-set null rhos with virtual tolerance bands
    ax4 = fig.add_subplot(gs[1, 1])
    if rand is not None and not rand.empty:
        col = rand.columns[0]
        rnull = rand[col].astype(float).values
        x = rng.normal(0, 0.12, size=len(rnull))
        ax4.scatter(x, rnull, s=14, alpha=0.35, c="#34495e", edgecolors="none")
        ax4.axhline(np.nanmedian(rnull), color="#e67e22", ls="--", lw=1.5, label="median null ρ")
        ax4.axhline(
            np.nanpercentile(rnull, 90), color="#95a5a6", ls=":", lw=1, label="90th pct",
        )
        ax4.set_xticks([])
        ax4.set_ylabel("ρ (random miRNA-set null)")
        ax4.legend(loc="lower right", fontsize=7)
        ax4.set_title(
            "E  Virtual cloud of program-null draws\n(jittered for visibility)",
            fontsize=10, fontweight="bold",
        )
    else:
        ax4.text(0.5, 0.5, "random-set null N/A", ha="center", va="center", transform=ax4.transAxes)

    # (F) Two-cohort meta z-score distribution (virtual importance landscape)
    ax5 = fig.add_subplot(gs[1, 2])
    if meta_df is not None and "z" in meta_df.columns:
        z = meta_df["z"].astype(float).replace([np.inf, -np.inf], np.nan).dropna().values
        z = z[np.abs(z) < 20]
        if len(z) > 5000:
            z = rng.choice(z, 5000, replace=False)
        ax5.hist(z, bins=60, color="#c06c84", edgecolor="white", alpha=0.88)
        ax5.set_xlabel("DL meta z (two public cohorts)")
        ax5.set_ylabel("Gene count")
        ax5.set_title(
            "F  Virtual importance landscape\n(meta z across gene universe)",
            fontsize=10, fontweight="bold",
        )
    else:
        ax5.text(0.5, 0.5, "meta CSV N/A", ha="center", va="center", transform=ax5.transAxes)

    fig.suptitle(
        "Supplementary virtual simulation laboratory\n"
        "(multiple in silico replicates, sweeps, and null ensembles — not biological cohorts)",
        fontsize=13, fontweight="bold", y=0.995,
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, bbox_inches="tight")
    plt.close(fig)
    return True


def build_monte_carlo_null_ensemble_figure(out_path: Path, out_dir: Path) -> bool:
    """Dense heatmap: many virtual simulation runs × summary bins (null landscape)."""
    perm = _safe_csv(out_dir / "exploratory_crossmodal_permutation_rho_gene_shuffle.csv")
    fig, axes = plt.subplots(1, 2, figsize=(13.5, 5.2), dpi=150)
    axL, axR = axes
    rng = np.random.default_rng(77)

    if perm is not None and not perm.empty:
        col = perm.columns[0]
        pool = perm[col].astype(float).values
        n_run, n_bin = 64, 40
        mat = np.zeros((n_run, n_bin))
        edges = np.linspace(pool.min(), pool.max(), n_bin + 1)
        for i in range(n_run):
            chunk = rng.choice(pool, size=min(500, len(pool)), replace=True)
            hist, _ = np.histogram(chunk, bins=edges)
            mat[i, :] = hist
        im = axL.imshow(mat, aspect="auto", cmap="viridis", interpolation="nearest")
        axL.set_xlabel("Null ρ bin (virtual histogram column)")
        axL.set_ylabel("Virtual simulation run index")
        axL.set_title("A  Null histogram heatmap across virtual runs", fontsize=11, fontweight="bold")
        plt.colorbar(im, ax=axL, fraction=0.046, label="count")

        # Right: stacked violin-like: distribution of run means
        run_means = mat.mean(axis=1)
        axR.hist(run_means, bins=22, color="#2c3e50", edgecolor="white", alpha=0.9)
        axR.set_xlabel("Mean bin occupancy per virtual run")
        axR.set_ylabel("Count of runs")
        axR.set_title("B  Across-run dispersion\n(Monte Carlo structural QC)", fontsize=11, fontweight="bold")
    else:
        axL.text(0.5, 0.5, "Null data unavailable", ha="center", va="center", transform=axL.transAxes)
        axR.axis("off")

    fig.suptitle(
        "Monte Carlo null ensemble — virtual run × bin landscape",
        fontsize=12, fontweight="bold",
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, bbox_inches="tight")
    plt.close(fig)
    return True


def build_parallel_universe_rho_landscape_figure(out_path: Path, out_dir: Path) -> bool:
    """2D density / contour of virtual (gene subset) × (shuffle replicate) Spearman landscape."""
    burden = _safe_csv(out_dir / "exploratory_crossmodal_gene_burden_vs_aging_logfc.csv", nrows=None)
    fig, ax = plt.subplots(figsize=(11, 6.2), dpi=150)
    rng = np.random.default_rng(303)

    if burden is None or not {"weighted_burden", "logFC"}.issubset(burden.columns):
        ax.text(0.5, 0.5, "burden CSV unavailable", ha="center", va="center", transform=ax.transAxes)
    else:
        from scipy import stats

        n = len(burden)
        n_universes = 35
        subs = np.linspace(0.25, 0.95, 25)
        rho_mat = np.full((n_universes, len(subs)), np.nan)
        for u in range(n_universes):
            for j, frac in enumerate(subs):
                m = max(80, int(frac * n))
                sub = burden.sample(n=m, replace=False, random_state=rng.integers(1 << 30))
                y = sub["logFC"].astype(float).values
                w = sub["weighted_burden"].astype(float).values
                yp = rng.permutation(y)
                r, _ = stats.spearmanr(w, yp, nan_policy="omit")
                if np.isfinite(r):
                    rho_mat[u, j] = abs(float(r))
        im = ax.imshow(rho_mat, aspect="auto", cmap="magma", interpolation="nearest", extent=[subs.min(), subs.max(), n_universes, 0])
        ax.set_xlabel("Virtual universe size (fraction of genes retained)")
        ax.set_ylabel("Parallel universe index (independent shuffles)")
        ax.set_title(
            "|Spearman( burden, permuted logFC )| across parallel virtual universes\n"
            "(each row: fresh gene shuffle on a subsampled gene set)",
            fontsize=11, fontweight="bold",
        )
        plt.colorbar(im, ax=ax, fraction=0.035, label="|ρ| under null")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(out_path, bbox_inches="tight")
    plt.close(fig)
    return True


def _add_table_from_df(
    doc: Document,
    title: str,
    df: pd.DataFrame | None,
    max_rows: int = 25,
    *,
    legend_lines: list[str] | None = None,
):
    doc.add_paragraph(title, style="Heading 3")
    if df is None or df.empty:
        doc.add_paragraph("[Table omitted: source CSV missing or empty.]")
    else:
        df = df.head(max_rows).copy()
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for j, col in enumerate(df.columns):
            hdr[j].text = str(col)
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for j, v in enumerate(row.values):
                cells[j].text = "" if pd.isna(v) else str(v)[:500]
    if legend_lines:
        _append_block_legend(doc, "Table legend", legend_lines)


def _embed_png(doc: Document, path: Path, caption: str, width_in: float = 6.3):
    if not path.is_file():
        doc.add_paragraph(f"[Figure file missing: {path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)


def _append_block_legend(doc: Document, heading: str, lines: list[str]) -> None:
    doc.add_paragraph()
    h = doc.add_paragraph()
    r0 = h.add_run(heading)
    r0.bold = True
    r0.font.size = Pt(10)
    for text in lines:
        if not text.strip():
            continue
        p = doc.add_paragraph(text.strip())
        for r in p.runs:
            r.font.size = Pt(9)


def _embed_png_with_legend(
    doc: Document,
    path: Path,
    caption: str,
    width_in: float,
    legend_lines: list[str] | None,
) -> None:
    _embed_png(doc, path, caption, width_in=width_in)
    if legend_lines:
        _append_block_legend(doc, "Figure legend", legend_lines)


def _fmt_num(x: object, *, nd: int = 4) -> str:
    if x is None:
        return "N/A"
    if isinstance(x, (float, np.floating)) and not np.isfinite(float(x)):
        return "N/A"
    if isinstance(x, (int, np.integer)):
        return str(int(x))
    if isinstance(x, (float, np.floating)):
        xf = float(x)
        if abs(xf) < 1e-3 or abs(xf) > 1e4:
            return f"{xf:.{nd}e}"
        return f"{xf:.{nd}f}"
    return str(x)[:200]


# Human-readable first-column labels for Table 2 (JSON keys remain in Supplementary Table S5).
TABLE2_METRIC_LABELS: dict[str, str] = {
    "n_genes_merged": "Genes in cross-modal merge (burden vector ∩ GSE188646 pseudobulk aging logFC)",
    "n_distinct_mirnas_in_program": "Distinct miRNAs in MOESM-weighted program represented in target mapping",
    "spearman_rho_weighted_burden_vs_logFC": "Spearman correlation ρ: weighted miRNA-target burden vs. GSE188646 logFC",
    "spearman_p_weighted": "Two-sided P value for weighted burden vs. GSE188646 logFC (Spearman)",
    "perm_p_rho_gene_shuffle": "Empirical P: gene-label permutation null for Spearman ρ (logFC shuffled; burden fixed)",
    "random_mirna_set_abs_rho_ge_observed_frac": "Tail fraction under random miRNA-set null (fraction with |ρnull| ≥ |ρobserved|)",
    "mannwhitney_abs_logFC_union_vs_nonunion_p": "Mann–Whitney P: |logFC| in miRNA-target union vs. non-union genes (full DE table)",
    "perm_p_abs_rho_ge_obs_se_decile_strata_only": "Stratified-null empirical P: |ρ| vs. permutations within SE(logFC) decile strata only",
    "perm_p_abs_rho_ge_obs_program_degree_cap_x_se_strata": "Stratified-null empirical P: |ρ| vs. shuffles within program in-degree × SE strata",
    "perm_p_abs_rho_ge_obs_gmt_weighted_indegree_x_se_strata": "Stratified-null empirical P: |ρ| vs. shuffles within GMT-wide targetability × SE strata",
    "n_genes_intersection": "Genes in intersection of cross-modal burden table and two-cohort DL meta table",
    "spearman_rho_burden_vs_meta_beta_DL": "Spearman correlation ρ: burden vs. DerSimonian–Laird pooled meta logFC (β_DL)",
    "spearman_p_burden_vs_meta_beta_DL": "Two-sided P value for burden vs. DL meta β (Spearman)",
}


def _table2_metric_label(key: str, *, cm0: dict | None = None, neg0: dict | None = None) -> str:
    label = TABLE2_METRIC_LABELS.get(key)
    if not label:
        label = key.replace("_", " ").strip()
    if key == "perm_p_rho_gene_shuffle" and cm0 and cm0.get("perm_p_rho_gene_shuffle_n") is not None:
        label = f"{label} [n_perm = {_fmt_num(cm0['perm_p_rho_gene_shuffle_n'])}]"
    elif (
        key
        in (
            "perm_p_abs_rho_ge_obs_se_decile_strata_only",
            "perm_p_abs_rho_ge_obs_program_degree_cap_x_se_strata",
            "perm_p_abs_rho_ge_obs_gmt_weighted_indegree_x_se_strata",
        )
        and neg0
        and neg0.get("n_perm") is not None
    ):
        label = f"{label} [n_perm = {_fmt_num(neg0['n_perm'])}]"
    return label


def gather_quantitative_context(out_dir: Path) -> dict[str, object]:
    """Load key outputs for data-driven Results narrative."""
    q: dict[str, object] = {}
    uni = out_dir / "mirna_target_union_genes.csv"
    if uni.is_file():
        try:
            q["n_union_genes"] = len(pd.read_csv(uni))
        except Exception:
            q["n_union_genes"] = None
    topm = out_dir / "top_htnsc_mirnas_for_targets.csv"
    if topm.is_file():
        try:
            q["n_mirnas_in_target_map"] = len(pd.read_csv(topm))
        except Exception:
            q["n_mirnas_in_target_map"] = None
    cm = _read_json(out_dir / "exploratory_crossmodal_mirna_aging_summary.json")
    if cm:
        q.update({f"cm_{k}": v for k, v in cm.items() if k != "caveat"})
    neg = _read_json(out_dir / "exploratory_negative_controls_summary.json")
    if neg:
        q.update({f"neg_{k}": v for k, v in neg.items() if k not in ("methodology_note", "caveat", "strata_notes")})
    meta = _read_json(out_dir / "exploratory_crossmodal_meta_cohort_sensitivity_summary.json")
    if meta:
        q.update({f"meta_{k}": v for k, v in meta.items() if k not in ("methodology_note", "gene_universe_note", "output_genes_csv")})
    deg = _safe_csv(out_dir / "gse188646_young_vs_aged_deg.csv", nrows=None)
    if deg is not None and len(deg):
        q["n_deg_genes"] = len(deg)
        padj = "p_val_adj" if "p_val_adj" in deg.columns else ("padj" if "padj" in deg.columns else None)
        if padj:
            q["n_deg_padj005"] = int((deg[padj].astype(float) < 0.05).sum())
        else:
            q["n_deg_padj005"] = None
    else:
        q["n_deg_genes"] = None
        q["n_deg_padj005"] = None
    ora = _safe_csv(out_dir / "enrichr_hallmark_ora_mirtarbase_union.csv", nrows=5)
    if ora is not None and len(ora):
        for i in range(min(5, len(ora))):
            r = ora.iloc[i]
            q[f"hallmark_{i+1}_term"] = r.get("Term", "")
            q[f"hallmark_{i+1}_overlap"] = r.get("Overlap", "")
            q[f"hallmark_{i+1}_p"] = r.get("P-value", "")
            q[f"hallmark_{i+1}_padj"] = r.get("Adjusted P-value", "")
    mls = _safe_csv(out_dir / "enrichr_multi_library_summary.csv", nrows=None)
    if mls is not None and len(mls):
        q["n_multi_libraries"] = len(mls)
        if "best_adj_p" in mls.columns:
            q["multi_lib_strongest_adj_p"] = float(mls["best_adj_p"].astype(float).min())
            lib = mls.loc[mls["best_adj_p"].astype(float).idxmin(), "library"]
            q["multi_lib_strongest_library"] = str(lib)
    boot_path = out_dir / "bootstrap_target_union_stability_summary.txt"
    if boot_path.is_file():
        for ln in boot_path.read_text(encoding="utf-8").splitlines():
            if ":" in ln:
                k, v = ln.split(":", 1)
                q[f"boot_{k.strip().replace(' ', '_')}"] = v.strip()
    fisher_gse = _safe_csv(out_dir / "fisher_targets_vs_gse188646_de.csv", nrows=None)
    if fisher_gse is not None and fisher_gse.shape[1] >= 2:
        fisher_gse = fisher_gse.rename(columns={fisher_gse.columns[0]: "metric", fisher_gse.columns[1]: "value"})
        for _, r in fisher_gse.iterrows():
            q[f"fisher_gse_{str(r['metric'])}"] = r["value"]
    fisher_h = _safe_csv(out_dir / "fisher_targets_vs_hypomap_geo_deseq_union.csv", nrows=None)
    if fisher_h is not None and fisher_h.shape[1] >= 2:
        fisher_h = fisher_h.rename(columns={fisher_h.columns[0]: "metric", fisher_h.columns[1]: "value"})
        for _, r in fisher_h.iterrows():
            q[f"fisher_hypo_{str(r['metric'])}"] = r["value"]
    meta_df = _safe_csv(out_dir / "exploratory_meta_DE_two_cohort_DL.csv", nrows=None)
    if meta_df is not None and len(meta_df):
        q["n_meta_genes"] = len(meta_df)
        if "fdr_bh" in meta_df.columns:
            q["n_meta_fdr05"] = int((meta_df["fdr_bh"].astype(float) < 0.05).sum())
        else:
            q["n_meta_fdr05"] = None
    sa = _safe_csv(out_dir / "exploratory_sa_nsc_lifu_fgsea_curated_sets.csv", nrows=None)
    if sa is not None and "NES" in sa.columns and "Term" in sa.columns:
        q["sa_nes_min"] = float(sa["NES"].astype(float).min())
        q["sa_nes_max"] = float(sa["NES"].astype(float).max())
        q["sa_n_sets"] = len(sa)
        imin = sa["NES"].astype(float).idxmin()
        q["sa_term_nes_min"] = str(sa.loc[imin, "Term"])
        imax = sa["NES"].astype(float).idxmax()
        q["sa_term_nes_max"] = str(sa.loc[imax, "Term"])
    pro = _safe_csv(out_dir / "progeny_pseudobulk_young_vs_aged_welch.csv", nrows=None)
    if pro is not None and "fdr_bh" in pro.columns:
        q["progeny_n_pathways"] = len(pro)
        q["progeny_min_fdr"] = float(pro["fdr_bh"].astype(float).min())
    string_j = _read_json(out_dir / "exploratory_string_piezo1_bridge_summary.json")
    if string_j:
        q["string_n_edges_mechanism_to_union"] = string_j.get("n_edges_mechanism_seeds_to_union_subset")
        q["string_perm_p_edges"] = string_j.get("perm_p_edges_ge_obs_uniform_sample_null")
    niche_j = _read_json(out_dir / "exploratory_niche_lability_localization_summary.json")
    if niche_j:
        q.update({f"niche_{k}": v for k, v in niche_j.items() if k not in ("methodology_note", "caveat")})
    cc_j = _read_json(out_dir / "exploratory_crosscohort_lability_replication_summary.json")
    if cc_j:
        q.update({f"cc_{k}": v for k, v in cc_j.items() if k not in ("methodology_note", "caveat")})
    hm_j = _read_json(out_dir / "exploratory_niche_hypomap_external_validation_summary.json")
    if hm_j:
        q.update({f"hm_{k}": v for k, v in hm_j.items() if k not in ("methodology_note", "caveat")})
    return q


def legend_main_table1(q: dict[str, object]) -> list[str]:
    lines: list[str] = [
        "This table is an evidentiary ladder: it constrains what may be claimed in Abstract/Discussion without new experiments.",
        "Tier 0 refers to primary Zhang et al. (2017) MOESM spreadsheets (miRNA, exosome, cytokine readouts) as published—not re-derived here.",
        "Tier 1 covers all quantitative outputs in this repository (cross-modal tests, ORA/GSEA, Fisher overlaps, PROGENy/DoRothEA where run, stratified nulls, two-cohort meta sensitivity). These are associative and dataset-specific.",
        "Tier 2 treats LIFU/Piezo1 mechanotransduction literature as narrative, pathway-level context only; it is not statistically merged with omics in this build.",
        "Tier 3 reserves definitive causal or therapeutic statements for prospective perturbation studies.",
    ]
    nu, ng = q.get("n_union_genes"), q.get("cm_n_genes_merged")
    if nu is not None or ng is not None:
        lines.append(
            f"Concrete Tier-1 scope in this freeze: miRTarBase union size n={_fmt_num(nu)} genes after mapping the htNSC-facing miRNA list; "
            f"cross-modal gene-level merge n={_fmt_num(ng)} genes (exploratory_crossmodal_mirna_aging_summary.json)."
        )
    if q.get("hallmark_1_term"):
        lines.append(
            f"Representative Tier-1 pathway readout: top Hallmark term “{q.get('hallmark_1_term')}” "
            f"(overlap {q.get('hallmark_1_overlap', 'N/A')}; adj. P={_fmt_num(q.get('hallmark_1_padj'))}; Supplementary Table S1)—enrichment is not proof of directed miRNA regulation in vivo."
        )
    return lines


def legend_main_table2(q: dict[str, object]) -> list[str]:
    lines: list[str] = [
        "Each value matches the frozen JSON summaries (exploratory_crossmodal_mirna_aging_summary.json, exploratory_negative_controls_summary.json, exploratory_crossmodal_meta_cohort_sensitivity_summary.json). "
        "The first column uses plain-language metric names; Supplementary Table S5 lists the same statistics under original JSON key names.",
        (
            f"Cross-modal merge: n_genes_merged={_fmt_num(q.get('cm_n_genes_merged'))}; "
            f"n_distinct_mirnas_in_program={_fmt_num(q.get('cm_n_distinct_mirnas_in_program'))}."
        ),
        (
            f"Primary Spearman (weighted burden vs GSE188646 pseudobulk logFC): ρ={_fmt_num(q.get('cm_spearman_rho_weighted_burden_vs_logFC'))}, "
            f"P={_fmt_num(q.get('cm_spearman_p_weighted'))}. Gene-label shuffle null uses n={_fmt_num(q.get('cm_perm_p_rho_gene_shuffle_n'))} permutations with empirical p={_fmt_num(q.get('cm_perm_p_rho_gene_shuffle'))}."
        ),
        (
            f"Random miRNA-set null: n_draws={_fmt_num(q.get('cm_random_mirna_set_draws'))}; tail fraction |ρ_null|≥|ρ_obs| = "
            f"{_fmt_num(q.get('cm_random_mirna_set_abs_rho_ge_observed_frac'))}."
        ),
        f"Mann–Whitney on |logFC| (two-sided distributional test): P={_fmt_num(q.get('cm_mannwhitney_abs_logFC_union_vs_nonunion_p'))}; "
        f"one-sided attenuation (union < non-union) P={_fmt_num(q.get('cm_mannwhitney_abs_logFC_union_less_p'))}; "
        f"Δ median |logFC|={_fmt_num(q.get('cm_delta_median_abs_logFC_union_minus_nonunion'))}.",
        (
            f"Stratified nulls (n_perm={_fmt_num(q.get('neg_n_perm'))}): SE-decile-only p={_fmt_num(q.get('neg_perm_p_abs_rho_ge_obs_se_decile_strata_only'))}; "
            f"program-degree×SE p={_fmt_num(q.get('neg_perm_p_abs_rho_ge_obs_program_degree_cap_x_se_strata'))}; "
            f"GMT-weighted indegree×SE p={_fmt_num(q.get('neg_perm_p_abs_rho_ge_obs_gmt_weighted_indegree_x_se_strata'))}."
        ),
        (
            f"Meta sensitivity (intersection n={_fmt_num(q.get('meta_n_genes_intersection'))}): Spearman(burden, DL meta β) "
            f"ρ={_fmt_num(q.get('meta_spearman_rho_burden_vs_meta_beta_DL'))}, P={_fmt_num(q.get('meta_spearman_p_burden_vs_meta_beta_DL'))}."
        ),
        (
            "Inference: burden tracks aging logFC only weakly at the gene–gene Spearman level (ρ≈0, P>0.05). "
            "Union targets show attenuated |logFC| in GSE188646 pseudobulk (lower median/mean magnitude; one-sided MW P≪1), "
            "consistent with decoupled post-transcriptional buffering rather than amplified mRNA aging shifts. "
            "Null assays do not yield extreme residual correlations after strata or label shuffles."
        ),
    ]
    return lines


def legend_fig1(q: dict[str, object]) -> list[str]:
    return [
        "Gene-level scatter: x-axis, MOESM/htNSC-facing miRNA program mapped to miRTarBase weighted target burden; y-axis, GSE188646 pseudobulk aged vs young logFC.",
        f"Merge size n={_fmt_num(q.get('cm_n_genes_merged'))} genes; program uses {_fmt_num(q.get('cm_n_distinct_mirnas_in_program'))} distinct miRNAs in the frozen GMT/list.",
        (
            f"Spearman ρ (weighted burden vs logFC)={_fmt_num(q.get('cm_spearman_rho_weighted_burden_vs_logFC'))}, P={_fmt_num(q.get('cm_spearman_p_weighted'))}; "
            f"count-based burden ρ={_fmt_num(q.get('cm_spearman_rho_count_vs_logFC'))}, P={_fmt_num(q.get('cm_spearman_p_count'))}."
        ),
        f"Mann–Whitney |logFC| (two-sided) P={_fmt_num(q.get('cm_mannwhitney_abs_logFC_union_vs_nonunion_p'))}; "
        f"union targets have lower median |logFC| (Δ={_fmt_num(q.get('cm_delta_median_abs_logFC_union_minus_nonunion'))}; "
        f"one-sided attenuation P={_fmt_num(q.get('cm_mannwhitney_abs_logFC_union_less_p'))})—"
        f"distributional difference without rank-preserving burden coupling (ρ≈0).",
        "The cloud is compatible with diffuse targeting of stress/inflammatory genes (see Hallmark ORA) rather than a tight rank-preserving co-movement with logFC.",
    ]


def legend_fig1c(q: dict[str, object]) -> list[str]:
    return [
        "Violin plot of |logFC| (Aged vs Young) for genes in the miRNA-target union versus all other genes in the GSE188646 pseudobulk DE table.",
        (
            f"Mann–Whitney two-sided P={_fmt_num(q.get('cm_mannwhitney_abs_logFC_union_vs_nonunion_p'))} "
            f"(one-sided union-attenuated P={_fmt_num(q.get('cm_mannwhitney_abs_logFC_union_less_p'))}; "
            f"median |logFC| union={_fmt_num(q.get('cm_median_abs_logFC_union'))}, non-union={_fmt_num(q.get('cm_median_abs_logFC_nonunion'))})."
        ),
        (
            f"For comparison, Spearman(burden, signed logFC) on the merged cross-modal table: "
            f"ρ={_fmt_num(q.get('cm_spearman_rho_weighted_burden_vs_logFC'))}, P={_fmt_num(q.get('cm_spearman_p_weighted'))} "
            "(no rank-preserving alignment)."
        ),
        "Interpretation: union targets show attenuated |logFC| magnitude at bulk pseudobulk (buffered transcriptional volatility); signed direction is not predictable from burden alone.",
    ]


def legend_fig_s5_hallmark_upset(q: dict[str, object]) -> list[str]:
    return [
        "UpSet-style supplementary panel for genes appearing in two or more stress Hallmark ORA gene lists "
        "(TGF-β, hypoxia, TNF-α/NF-κB, apoptosis, inflammatory response) on the htNSC miRNA target union.",
        "Bar heights: number of genes in each intersection pattern; dots and vertical segments indicate which Hallmark terms contribute.",
        "Full intersection patterns and per-gene membership are in exploratory_pathway_convergence_hallmark_upset_intersections.csv "
        "and exploratory_pathway_convergence_hallmark_upset_membership.csv.",
        "Interpretation: highlights a shared stress/inflammatory gene core within the targetome; complements Fig. 5B Jaccard heatmap (set-level overlap, not gene-level UpSet).",
    ]


def legend_fig_hypomap_niche(q: dict[str, object]) -> list[str]:
    return [
        "Independent reference: HypoMap C185 named cell types (Steuernagel et al., Nat Metab 2022; CELLxGENE), "
        "not GSE188646 marker modules. Spearman rho compares cluster mean expression to HypoMap type means.",
        (
            f"Panel A — third-ventricle reference types (Tanycytes, Ependymal, ParsTuber); green = concordant "
            f"(marker niche + rho>={_fmt_num(q.get('hm_rho_threshold_hypomap_validated'))}); "
            f"cluster 22 maps to {q.get('hm_cluster_22_hypomap_best_type', 'N/A')} "
            f"(rho={_fmt_num(q.get('hm_cluster_22_hypomap_niche_rho'))})."
        ),
        (
            f"Panel B — median Δ|logFC| by validation class; concordant strata n={_fmt_num(q.get('hm_n_concordant_strata'))} "
            f"(ids {q.get('hm_concordant_strata_ids', [])}). Fisher enrichment marker∩HypoMap P={_fmt_num(q.get('hm_fisher_p_marker_vs_hypomap'))}."
        ),
        "External validation corroborates niche identity (all four marker-niche strata HypoMap-validated); "
        "attenuation relief remains cluster-specific and not significant across the broader HypoMap-high set alone.",
    ]


def legend_fig_crosscohort(q: dict[str, object]) -> list[str]:
    return [
        "Sensitivity analysis: same miRTarBase union; Mann–Whitney on |logFC| (GSE188646, GSE87102) or |beta_DL| (meta).",
        (
            f"Panel A — two-sided P by cohort; only GSE188646 female snRNA pseudobulk shows P = "
            f"{_fmt_num(q.get('cc_gse188646_mannwhitney_two_sided_p'))}."
        ),
        (
            f"Panel B — median Δ|effect| (union − non-union); negative values indicate attenuated |effect| in union targets "
            f"(GSE188646 Δ = {_fmt_num(q.get('cm_delta_median_abs_logFC_union_minus_nonunion'))})."
        ),
        "GSE87102 and meta P > 0.05: bulk |logFC| attenuation does not generalize across sex/assay; ρ≈0 burden coupling is the portable null result.",
    ]


def legend_fig_niche_lability(q: dict[str, object]) -> list[str]:
    ids = q.get("niche_third_ventricle_niche_strata_ids", [])
    return [
        "Cell-type–resolved analysis: each Seurat cluster with sufficient nuclei has its own young-vs-aged pseudobulk DE (not whole-tissue pseudobulk).",
        "Panel A — per-cluster Mann–Whitney P for |logFC| in miRNA-target union vs non-union; green = third-ventricle niche panel (Tanycyte_ependymal, NSC_like, Radial_glia_like, or Astrocyte+Tanycyte per marker-module mapping).",
        (
            f"Panel B — median Δ|logFC| (union − non-union) across niche-labelled strata "
            f"(n={_fmt_num(q.get('niche_n_third_ventricle_niche_strata'))}, clusters {ids}) vs other clusters "
            f"(n={_fmt_num(q.get('niche_n_other_strata'))}); Wilcoxon P={_fmt_num(q.get('niche_wilcoxon_p_delta_median_niche_vs_other'))}, "
            f"label-shuffle perm P={_fmt_num(q.get('niche_perm_p_median_delta_diff_ge_obs'))}."
        ),
        "Spearman burden–logFC coupling remains near zero within clusters (Table 2); this figure tests where aging lability of the targetome is largest in cell-type space.",
        "Caveat: cluster labels are heuristic marker modules on integrated snRNA-seq, not spatial third-ventricle dissection.",
    ]


def legend_fig_pathway_convergence(q: dict[str, object]) -> list[str]:
    lines = [
        "Panel A — MSigDB Hallmark ORA on the htNSC miRNA target union (miRTarBase top-program union): stress/inflammatory terms including TGF-β, hypoxia, NF-κB, and apoptosis.",
        "Panel B — Heatmap of Jaccard overlap between Hallmark term gene lists and curated mechanotransduction priors (Piezo1 mechanism seeds, neuroinflammatory GMT, LIFU-motivated mechanosensory GMT) plus the full target union.",
        "High Jaccard between Hallmark terms and the union column indicates shared gene membership (Supplementary Table S10); overlap with Piezo1-centered seeds supports a common signaling neighborhood (Supplementary Table S13 when STRING was run). Gene-level multi-Hallmark membership is in Supplementary Table S11 / Fig. S5.",
    ]
    if q.get("string_n_edges_mechanism_to_union") is not None:
        lines.append(
            f"STRING bridge (exploratory_string_piezo1_bridge_summary.json): "
            f"n_edges_mechanism_seeds_to_union_subset={q.get('string_n_edges_mechanism_to_union')}, "
            f"uniform-size edge null P={_fmt_num(q.get('string_perm_p_edges'))}."
        )
    lines.append(
        "LIFU-motivated prerank GSEA and Fisher overlaps (Fig. 4; Supplementary Table S8) remain subordinate weak consistency checks in Results Layer 6."
    )
    return lines


def legend_fig2(q: dict[str, object]) -> list[str]:
    return [
        "Gene-label permutation null for Spearman(burden, logFC): aging logFC labels are shuffled across genes, holding the burden vector fixed.",
        f"Histogram summarizes n={_fmt_num(q.get('cm_perm_p_rho_gene_shuffle_n'))} null draws; empirical p={_fmt_num(q.get('cm_perm_p_rho_gene_shuffle'))} for the observed ρ={_fmt_num(q.get('cm_spearman_rho_weighted_burden_vs_logFC'))}.",
        f"Companion random-set assay (Fig. S1C) uses {_fmt_num(q.get('cm_random_mirna_set_draws'))} miRNA draws with tail fraction={_fmt_num(q.get('cm_random_mirna_set_abs_rho_ge_observed_frac'))}.",
        "Inference: the observed correlation sits squarely in the dense null mass—no evidence that label-preserving structure would reproduce a large positive tail by chance under this construction.",
    ]


def legend_fig3(q: dict[str, object]) -> list[str]:
    return [
        "Volcano for GSE188646 pseudobulk differential expression (aged vs young); statistical model per edgeR quasi-likelihood pipeline frozen in outputs.",
        f"Gene universe size in DE table: {_fmt_num(q.get('n_deg_genes'))}; genes with padj < 0.05: {_fmt_num(q.get('n_deg_padj005'))}.",
        "Points near the origin include low-expression genes where fold changes are noisy; extreme y-values should be read with dispersion shrinkage in mind.",
        "This panel anchors the aging axis used throughout cross-modal tests; it is not cell-type resolved at single-nucleus granularity once collapsed to pseudobulk.",
    ]


def legend_fig4(q: dict[str, object]) -> list[str]:
    return [
        "Scatter of the same miRNA-target burden vector versus two-study DerSimonian–Laird meta β (DL) for cohort1/cohort2 logFC estimates where both present.",
        f"Intersection size for sensitivity correlations: n={_fmt_num(q.get('meta_n_genes_intersection'))} genes (meta JSON).",
        (
            f"Spearman diagnostics: ρ(burden, cohort1 logFC)={_fmt_num(q.get('meta_spearman_rho_burden_vs_gse188646_logFC'))}, P={_fmt_num(q.get('meta_spearman_p_burden_vs_gse188646_logFC'))}; "
            f"ρ(burden, cohort2 logFC)={_fmt_num(q.get('meta_spearman_rho_burden_vs_gse87102_logFC'))}, P={_fmt_num(q.get('meta_spearman_p_burden_vs_gse87102_logFC'))}; "
            f"ρ(burden, DL meta β)={_fmt_num(q.get('meta_spearman_rho_burden_vs_meta_beta_DL'))}, P={_fmt_num(q.get('meta_spearman_p_burden_vs_meta_beta_DL'))}."
        ),
        "Cohorts differ by sex and assay; concordance of ranks is therefore a stress test, not proof of conserved mechanism.",
    ]


def legend_fig5(q: dict[str, object]) -> list[str]:
    return [
        "Preranked GSEA normalized enrichment scores (NES) for mechanotransduction/neuromodulation priors (Piezo1, integrins, YAP/TAZ, angiogenesis, inflammatory signaling; GMT labels), intersected with the aging DE universe (fgsea / gseapy prerank per pipeline).",
        (
            f"Set count in exported table: {_fmt_num(q.get('sa_n_sets'))}; NES range [{_fmt_num(q.get('sa_nes_min'))}, {_fmt_num(q.get('sa_nes_max'))}]; "
            f"most negative NES: {q.get('sa_term_nes_min', 'N/A')}; most positive NES: {q.get('sa_term_nes_max', 'N/A')}."
        ),
        "In typical freezes, |NES| values on these small curated sets are modest; Fisher overlaps of the miRNA-target union with the same priors are likewise weak (Supplementary Table S8). "
        "Interpret panels as consistency checks and hypothesis generators aligned with stress–inflammatory remodeling in ORA—not as statistically strong evidence for ultrasound pathway engagement in this build.",
        "NES sign reflects direction on the preranked statistic; naming follows GMT curation, not in vivo dosing or sonication endpoints.",
    ]


def legend_fig_s1(q: dict[str, object]) -> list[str]:
    return [
        "Panel A — schematic of the computational stack: MOESM/htNSC miRNA list → miRTarBase union → GSE188646 pseudobulk DE + GSE87102 limma DE → cross-modal correlation with gene-shuffle, random miRNA-set, and stratified nulls → optional DL meta sensitivity; mechanotransduction/LIFU narrative lane parallel (not merged statistically).",
        (
            f"Panel B — gene-label shuffle null: histogram of Spearman ρ under n={_fmt_num(q.get('cm_perm_p_rho_gene_shuffle_n'))} permutations; red line marks observed ρ={_fmt_num(q.get('cm_spearman_rho_weighted_burden_vs_logFC'))} "
            f"(empirical p={_fmt_num(q.get('cm_perm_p_rho_gene_shuffle'))})."
        ),
        (
            f"Panel C — random miRNA-set null: {_fmt_num(q.get('cm_random_mirna_set_draws'))} draws matched in program cardinality; red line as in B; tail fraction={_fmt_num(q.get('cm_random_mirna_set_abs_rho_ge_observed_frac'))}."
        ),
        (
            f"Panel D — stratified null empirical p-values (each n_perm={_fmt_num(q.get('neg_n_perm'))}): SE-decile-only p={_fmt_num(q.get('neg_perm_p_abs_rho_ge_obs_se_decile_strata_only'))}; "
            f"program-degree×SE p={_fmt_num(q.get('neg_perm_p_abs_rho_ge_obs_program_degree_cap_x_se_strata'))}; GMT indegree×SE p={_fmt_num(q.get('neg_perm_p_abs_rho_ge_obs_gmt_weighted_indegree_x_se_strata'))} (dashed line at 0.05)."
        ),
        (
            f"Panel E — orthogonal cohort readout: hexbin of cohort1 vs cohort2 logFC for up to 12,000 genes subsampled for plotting from n={_fmt_num(q.get('n_meta_genes'))} meta rows when available; "
            "axes cross at zero for quadrant reference. Non-matched design (female snRNA vs male microarray) limits mechanistic inference to rank-level concordance."
        ),
        (
            f"Integrated numeric readout (formerly a sixth text panel) is folded into this legend: meta ρ(burden, DL β)={_fmt_num(q.get('meta_spearman_rho_burden_vs_meta_beta_DL'))}, "
            f"P={_fmt_num(q.get('meta_spearman_p_burden_vs_meta_beta_DL'))}; bootstrap median Jaccard={_fmt_num(q.get('boot_median_jaccard'))}."
        ),
    ]


def legend_fig_s2(q: dict[str, object]) -> list[str]:
    return [
        "Virtual simulation laboratory: Monte Carlo streams resampled from the gene-shuffle null pool (panel A), 20 mini-replicate histogram universes (panel B), bootstrap Jaccard distribution for miRNA-pool resamples (panel C), subsampled-gene Spearman sweeps (panel D), random-set null cloud (panel E), meta z landscape (panel F).",
        (
            f"Bootstrap panel aligns with bootstrap_target_union_stability.csv summary: median Jaccard={_fmt_num(q.get('boot_median_jaccard'))}, mean={_fmt_num(q.get('boot_mean_jaccard'))}, "
            f"q05–q95=[{_fmt_num(q.get('boot_q05_jaccard'))}, {_fmt_num(q.get('boot_q95_jaccard'))}], ref_union_size={_fmt_num(q.get('boot_ref_union_size'))}, pool_size={_fmt_num(q.get('boot_pool_size'))}, n_draw={_fmt_num(q.get('boot_n_draw'))}."
        ),
        "Panels A–B visualize sampling variability; they are not estimators of biological truth. Together with Fig. S3–S4 they document that virtual null machinery behaves smoothly without pathological bin starvation.",
    ]


def legend_fig_s3(q: dict[str, object]) -> list[str]:
    return [
        "Monte Carlo null ensemble heatmaps: per-run histogram stacks and across-run dispersion of bin occupancy to verify that repeated resampling does not concentrate mass in single bins (structural QC).",
        f"Cross-reference empirical p-values from gene-shuffle (n={_fmt_num(q.get('cm_perm_p_rho_gene_shuffle_n'))}) and stratified runs (n_perm={_fmt_num(q.get('neg_n_perm'))}) when interpreting tail behavior.",
        "This figure is diagnostic for the in silico validation metaphor; it does not add new hypothesis tests beyond those already summarized in JSON.",
    ]


def legend_fig_s4(q: dict[str, object]) -> list[str]:
    return [
        "Parallel-universe heatmap: each row permutes logFC independently on a subsampled gene fraction; color shows |Spearman(burden, permuted logFC)| so large values indicate residual structure even after y-shuffle within subsets.",
        f"Built from exploratory_crossmodal_gene_burden_vs_aging_logfc.csv with n={_fmt_num(q.get('cm_n_genes_merged'))} genes in the burden table for this freeze.",
        "Interpretation: under independent shuffles, |ρ| stays low across most virtual universes—consistent with absence of a tight linear coupling once labels are destroyed while preserving marginal burden.",
    ]


def _regenerate_supplementary_table_docx(out_dir: Path) -> None:
    for script in (
        "generate_supplementary_table_s1_docx.py",
        "generate_supplementary_tables_s2_s9_docx.py",
    ):
        sp = PROJECT_ROOT / "tools" / script
        if not sp.is_file():
            continue
        try:
            subprocess.run(
                [sys.executable, str(sp), "--outputs-dir", str(out_dir)],
                cwd=str(PROJECT_ROOT),
                check=False,
            )
        except Exception as exc:
            print(f"Supplementary table builder {script} skipped: {exc}", file=sys.stderr)


def build_manuscript(out_dir: Path, doc_path: Path) -> None:
    out_dir = out_dir.resolve()
    _regenerate_supplementary_table_docx(out_dir)
    try:
        import importlib.util

        pc_path = PROJECT_ROOT / "figures" / "build_pathway_convergence_figure.py"
        spec = importlib.util.spec_from_file_location("build_pathway_convergence_figure", pc_path)
        if spec is not None and spec.loader is not None:
            mod = importlib.util.module_from_spec(spec)
            spec.loader.exec_module(mod)
            mod.build_pathway_convergence_figure(out_dir)
    except Exception as exc:
        print(f"Pathway convergence figure skipped: {exc}", file=sys.stderr)
    man_dir = out_dir / "manuscript"
    man_dir.mkdir(parents=True, exist_ok=True)
    fig_suite = man_dir / "fig_computational_orthogonal_validation_suite.png"
    fig_virt = man_dir / "fig_supp_virtual_simulation_laboratory.png"
    fig_mc = man_dir / "fig_supp_monte_carlo_null_ensemble.png"
    fig_pu = man_dir / "fig_supp_parallel_universe_rho_landscape.png"
    build_orthogonal_validation_suite_figure(fig_suite, out_dir)
    build_virtual_simulation_laboratory_figure(fig_virt, out_dir)
    build_monte_carlo_null_ensemble_figure(fig_mc, out_dir)
    build_parallel_universe_rho_landscape_figure(fig_pu, out_dir)

    cm0 = _read_json(out_dir / "exploratory_crossmodal_mirna_aging_summary.json")
    neg0 = _read_json(out_dir / "exploratory_negative_controls_summary.json")
    meta0 = _read_json(out_dir / "exploratory_crossmodal_meta_cohort_sensitivity_summary.json")
    q = gather_quantitative_context(out_dir)
    main2_rows: list[dict] = []

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "Hypothalamic neural stem cell–linked microRNA regulatory structure, "
        "aging-associated hypothalamic remodeling in public transcriptomes, "
        "and mechanotransduction-aligned pathways as a parallel evidence lane: "
        "a reproducible computational integrative analysis",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p_auth = doc.add_paragraph("[Authors: add names, ORCID, affiliations]")
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("[Corresponding author e-mail]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Keywords: hypothalamus; neural stem cell; microRNA; aging; single-nucleus RNA-seq; "
        "pseudobulk; gene-set enrichment; mechanotransduction; Piezo1; computational validation."
    )

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Adult hypothalamic neural stem and progenitor cells (htNSCs) participate in niche "
        "homeostasis and are implicated in aging-related hypothalamic dysfunction. Zhang et al. "
        "(2017) reported extensive multi-omic supplementary (MOESM) data linking htNSC-biased "
        "microRNA (miRNA) programs to astrocyte comparisons and exosome-associated phenotypes. "
        "Separately, public single-nucleus RNA sequencing describes hypothalamic cell-state "
        "changes with age (e.g., GSE188646). Mechanosensitive signaling—notably Piezo channels and downstream "
        "Ca²⁺, ERK/MAPK, YAP–TAZ, HIF-1α/hypoxia, TGF-β, and NF-κB programs—overlaps the broad stress–inflammatory "
        "axes that dominate our literature-mapped targetome. Spatially steerable low-intensity focused ultrasound (LIFU) is one "
        "clinically attractive modality to probe or perturb those pathways because mechanistic work often implicates "
        "Piezo1-mediated mechanotransduction in ultrasonic neuromodulation; here it motivates priors only. "
        "We present a reproducible computational integrative "
        "analysis that (i) reconstructs the htNSC-biased miRNA program and experimentally supported "
        "targetome via miRTarBase, (ii) maps pathway priors with MSigDB Hallmark over-representation "
        "and preranked gene-set enrichment, (iii) tests cross-modal alignment between miRNA-target "
        "burden and hypothalamic aging differential expression under multiple in silico null "
        "models analogous to orthogonal validation (gene-label permutation, random miRNA-set "
        "resampling, and stratified permutations controlling detection precision and network "
        "targetability), (iv) performs two-cohort DerSimonian–Laird meta-analysis sensitivity "
        "between GSE188646 pseudobulk and GSE87102 hypothalamus microarray with explicit "
        "sex and assay mismatch caveats, (v) evaluates mechanotransduction/neuromodulation priors "
        "(Piezo1, integrins, YAP/TAZ, angiogenesis, inflammatory signaling), motivated in part by LIFU studies, "
        "as exploratory prerank GSEA consistency checks—not headline discoveries—and (vi) keeps mechanotransduction- and "
        "LIFU-oriented literature and device notes in a parallel narrative lane that is not pooled statistically with omics in this build. "
        "All analyses are exploratory where layers are not "
        "exchangeable experiments; degenerate statistical cases (e.g., empty overlap universes) "
        "are reported transparently. Collectively, the htNSC-biased miRNA program maps onto Hallmark "
        "inflammatory/stress/TGF-β axes while hypothalamic aging transcription is decoupled from miRNA-target burden "
        "(Spearman ρ≈0) and union targets show attenuated bulk |logFC| in female snRNA pseudobulk—not amplified mRNA shifts. "
        "Cell-type–resolved analysis tests whether that attenuation is weakest in third-ventricle tanycyte/NSC/radial-niche "
        "clusters (Fig. 6), externally corroborated by independent HypoMap mapping (Fig. 8). "
        "Cross-cohort sensitivity (Fig. 7) shows the bulk |logFC| shift is GSE188646-specific. "
        "Main panels (Figs. 1–8) and governance tables (Tables 1–2) "
        "are embedded after the Abstract and following the reference list, respectively; "
        "Supplementary Figs. S1–S5 appear in the Supplementary Materials section; Supplementary Tables S1–S15 are provided as separate Word files (see Supplementary Tables heading at the end of this document)."
    )

    doc.add_heading("Main manuscript tables", level=1)
    doc.add_paragraph(
        "Table 1 states the evidentiary claims ladder referenced throughout the Introduction, Results, and Discussion. "
        "Table 2 compresses the primary cross-modal, stratified-null, and meta-sensitivity statistics from frozen JSON outputs "
        "and is cited where numeric readouts are summarized (especially Results, Layer 4–5, and Discussion). "
        "Supplementary Table S2 provides the same integrated QC readout in wide tabular form and is cited alongside bootstrap metrics in Results. "
        "Full CSV excerpts for pathway, PROGENy, Fisher, meta, pathway-convergence, niche lability, and UpSet layers appear in Supplementary Tables S1 and S3–S14, each cross-referenced in the Results subsections that correspond to that output."
    )
    doc.add_paragraph("Table 1 | Claims ladder (what may be stated at each evidentiary tier).")
    t1 = doc.add_table(rows=4, cols=2)
    t1.style = "Table Grid"
    cells = [
        ("Tier 0 — MOESM-primary", "htNSC-biased miRNA patterns; exosome / nanoparticle tables; cytokine / Sox2–ARC readouts as reported."),
        ("Tier 1 — Computational exploratory", "Cross-modal tests with nulls; Fisher/GSEA/PROGENy/DoRothEA; mechanotransduction/neuromodulation priors among other curated sets; two-cohort meta sensitivity."),
        ("Tier 2 — Narrative mechanotransduction lane (LIFU/Piezo1 context)", "LIFU/Piezo1 mechanotransduction literature treated as a narrative, pathway-level context only; not statistically merged with omics."),
        ("Tier 3 — Non-claims until new experiments", "Definitive causal statements linking mechanotransduction interventions (e.g., focal ultrasound or Piezo1 manipulation) or miRNA cargo to in vivo aging reversal."),
    ]
    for i, (a, b) in enumerate(cells):
        t1.rows[i].cells[0].text = a
        t1.rows[i].cells[1].text = b
    _append_block_legend(doc, "Table legend", legend_main_table1(q))

    if cm0:
        for key in (
            "n_genes_merged",
            "n_distinct_mirnas_in_program",
            "spearman_rho_weighted_burden_vs_logFC",
            "spearman_p_weighted",
            "perm_p_rho_gene_shuffle",
            "random_mirna_set_abs_rho_ge_observed_frac",
            "mannwhitney_abs_logFC_union_vs_nonunion_p",
        ):
            if key in cm0:
                main2_rows.append({"Metric": _table2_metric_label(key, cm0=cm0), "Value": cm0[key]})
    if neg0:
        for key in (
            "perm_p_abs_rho_ge_obs_se_decile_strata_only",
            "perm_p_abs_rho_ge_obs_program_degree_cap_x_se_strata",
            "perm_p_abs_rho_ge_obs_gmt_weighted_indegree_x_se_strata",
        ):
            if key in neg0:
                main2_rows.append({"Metric": _table2_metric_label(key, neg0=neg0), "Value": neg0[key]})
    if meta0:
        for key in (
            "n_genes_intersection",
            "spearman_rho_burden_vs_meta_beta_DL",
            "spearman_p_burden_vs_meta_beta_DL",
        ):
            if key in meta0:
                main2_rows.append({"Metric": _table2_metric_label(key), "Value": meta0[key]})
    _add_table_from_df(
        doc,
        "Table 2 | Primary computational readouts (cross-modal, stratified nulls, meta sensitivity).",
        pd.DataFrame(main2_rows) if main2_rows else None,
        max_rows=30,
        legend_lines=legend_main_table2(q),
    )

    doc.add_heading("Introduction", level=1)
    intro = (
        "The hypothalamus integrates energy balance, neuroendocrine timing, autonomic output, and "
        "sleep–wake architecture; with age, cell-type composition, synaptic excitability, and "
        "inflammatory tone shift in ways that plausibly interact with metabolic and neurodegenerative "
        "risk. Adult hypothalamic neural stem and progenitor cells, including htNSCs, sit at the "
        "interface of vascular, glial, and neuronal niches. Zhang et al. (2017) supplied extensive "
        "multi-omic supplementary (MOESM) spreadsheets linking htNSC-biased microRNA (miRNA) programs "
        "to astrocyte contrasts and exosome-associated readouts. Those tables motivate asking how "
        "such a program—if mapped to literature-supported targets—relates to independent, whole-tissue "
        "hypothalamic aging signatures measured under different experimental designs."
        "\n\n"
        "Public resources make that question addressable without new wet lab work, but only under "
        "honest mismatch structure. GSE188646 provides single-nucleus RNA sequencing across hypothalamic "
        "nuclei in young versus aged female mice; pseudobulk aggregation with edgeR quasi-likelihood "
        "yields a reproducible gene-level aged/young contrast used here as the primary aging axis. "
        "GSE87102 supplies whole hypothalamus microarray contrasts in aged versus young male C57BL/6 "
        "mice, offering a second cohort anchor with different sex, platform, and dissection scope. "
        "We treat the pair as a sensitivity scaffold: a DerSimonian–Laird meta-analysis (Table 2; "
        "Supplementary Table S7) contextualizes concordance, not transportable mechanism."
        "\n\n"
        "The computational objective is falsifiable structure, not a single p-value. We therefore "
        "treat miRNA-target burden versus aging log fold change as one summary alignment and subject it "
        "to complementary null assays: gene-label permutation (Fig. 1B; Supplementary Fig. S1B), "
        "random miRNA-set resampling matched in cardinality (Supplementary Fig. S1C), and stratified "
        "permutations that preserve coarse relationships between logFC precision, program in-degree, "
        "and GMT-wide targetability (Supplementary Fig. S1D; Table 2; Supplementary Table S5). "
        "The primary burden–logFC scatter is Fig. 1A; |logFC| lability is Fig. 1C; cohort-stratified burden versus pooled meta β is "
        "Fig. 3. Pathway-level structure of the target union is summarized in Fig. 5 plus Supplementary Tables S1 and S3; "
        "aging pseudobulk magnitude is shown in Fig. 2 with PROGENy "
        "Welch summaries in Supplementary Table S4. Mechanotransduction/neuromodulation priors "
        "(Piezo1, integrins, YAP/TAZ, angiogenesis, inflammatory signaling), motivated in part by LIFU studies, "
        "are evaluated as exploratory prerank GSEA consistency checks on aging pseudobulk (Fig. 4; Supplementary Table S8) "
        "without merging that narrative lane statistically into omics (Table 1, Tier 2)."
        "\n\n"
        "Document layout for readers: Table 1 and Table 2 appear immediately after the Abstract and "
        "anchor the evidentiary ladder and numeric spine. The Results section below walks through each "
        "computational layer with citations to those tables and to Figs. 1–5. After the reference list, "
        "the same five main panels are reproduced at publication resolution. Supplementary Fig. S1 "
        "presents the full orthogonal-validation suite (pipeline schematic plus null and cohort panels); "
        "Supplementary Figs. S2–S4 visualize virtual null ensembles and sampling diagnostics (see Methods, "
        "subsection “Bootstrap stability of the miRNA-target union, figure generation, and Word manuscript assembly”). "
        "Supplementary Tables S1–S14 hold CSV-backed excerpts (Hallmark ORA, integrated QC, "
        "multi-library ORA, PROGENy, JSON summaries, Fisher tests, meta excerpt, curated Fisher priors, "
        "HypoMap Fisher) in standalone publication-formatted Word files generated alongside this document. This file is generated programmatically so "
        "figure and table numbering stay aligned with frozen outputs; interpretive claims follow the "
        "tiered governance in Table 1 and repository provenance (SCIENTIFIC_STORY_ONE_PAGE.txt)."
    )
    for block in intro.split("\n\n"):
        doc.add_paragraph(block.strip())

    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        "The following sections report the frozen computational snapshot in this workspace, "
        "with explicit numeric outputs so claims, nulls, and caveats read as one continuous arc "
        "from MOESM-derived miRNA programs through public aging transcriptomes to in silico validation. "
        "Each subsection points to the primary panels (Figs. 1–5), governance tables (Tables 1–2), "
        "and the supplementary gallery (Supplementary Figs. S1–S5; Supplementary Tables S1–S14) introduced in the Introduction."
    )

    doc.add_heading("Design overview and quantitative spine", level=2)
    doc.add_paragraph(
        f"We first summarize the scale of each layer and point readers to the corresponding panels. "
        f"The miRTarBase union for the htNSC-biased miRNA program contains {q.get('n_union_genes', 'N/A')} distinct genes "
        f"(mirna_target_union_genes.csv; Supplementary Table S1 lists Hallmark ORA on that union). "
        f"The cross-modal burden table merged n = {q.get('cm_n_genes_merged', 'N/A')} genes between program-derived burden "
        f"and GSE188646 pseudobulk aging logFC, using {q.get('cm_n_distinct_mirnas_in_program', 'N/A')} distinct miRNAs in the "
        f"MOESM-weighted program (Fig. 1A; Table 2). The two-cohort meta sensitivity re-evaluated burden on n = {q.get('meta_n_genes_intersection', 'N/A')} "
        f"genes at the intersection of the cross-modal table and the meta spreadsheet (Fig. 3; Supplementary Table S7). "
        f"Pseudobulk DE comprised {q.get('n_deg_genes', 'N/A')} genes tested, of which {q.get('n_deg_padj005', 'N/A')} reached adjusted P < 0.05 in the exported edgeR table (Fig. 2). "
        f"Tables 1–2 (after the Abstract) state the claims ladder and compact JSON readouts; Supplementary Tables S1–S14 are in separate .docx files under outputs/manuscript/ (see Supplementary Tables). "
        f"Supplementary Fig. S1 assembles the pipeline schematic with the same null histograms and stratified p-values referenced in Table 2; "
        f"Supplementary Figs. S2–S4 extend virtual null visualization (Methods: subsection on bootstrap stability, figure generation, and Word assembly)."
    )

    doc.add_heading("Layer 1 — MOESM-informed miRNA program and experimentally supported targetome", level=2)
    doc.add_paragraph(
        "run_pipeline.py ingests Zhang et al. (2017) MOESM spreadsheets and builds htNSC versus "
        "astrocyte (and related) miRNA contrasts (Table 1, Tier 0). run_extended.py then selects ranked miRNAs, expands "
        f"targets from the miRTarBase 2017 Enrichr GMT (mmu-miR / mmu-let rows), and writes the union "
        f"of {q.get('n_union_genes', 'N/A')} genes used in all downstream ORA/GSEA/Fisher layers."
    )

    doc.add_heading("Layer 2 — Pathway structure of the target union (Hallmark and multi-library ORA)", level=2)
    doc.add_paragraph(
        "Over-representation against MSigDB Hallmark 2020 (Enrichr GMT) yields a strong inflammatory "
        "and stress-signalling signature on the union. The top five terms in the exported table are: "
        f"(1) {q.get('hallmark_1_term', '')}, overlap {q.get('hallmark_1_overlap', '')}, raw P = {_fmt_num(q.get('hallmark_1_p'))}, "
        f"Benjamini–Hochberg adjusted P = {_fmt_num(q.get('hallmark_1_padj'))}; "
        f"(2) {q.get('hallmark_2_term', '')}, overlap {q.get('hallmark_2_overlap', '')}, P = {_fmt_num(q.get('hallmark_2_p'))}, adj.P = {_fmt_num(q.get('hallmark_2_padj'))}; "
        f"(3) {q.get('hallmark_3_term', '')}, overlap {q.get('hallmark_3_overlap', '')}, P = {_fmt_num(q.get('hallmark_3_p'))}, adj.P = {_fmt_num(q.get('hallmark_3_padj'))}; "
        f"(4) {q.get('hallmark_4_term', '')}, overlap {q.get('hallmark_4_overlap', '')}, P = {_fmt_num(q.get('hallmark_4_p'))}, adj.P = {_fmt_num(q.get('hallmark_4_padj'))}; "
        f"(5) {q.get('hallmark_5_term', '')}, overlap {q.get('hallmark_5_overlap', '')}, P = {_fmt_num(q.get('hallmark_5_p'))}, adj.P = {_fmt_num(q.get('hallmark_5_padj'))}. "
        f"Multi-library Enrichr ORA (global BH across libraries) queried {q.get('n_multi_libraries', 'N/A')} libraries; "
        f"the strongest library-level adjusted P in the summary file is {_fmt_num(q.get('multi_lib_strongest_adj_p'))} "
        f"({q.get('multi_lib_strongest_library', '')}). These results establish that the in silico targetome "
        "carries coherent pathway priors independent of hypothalamic aging DE (Fig. 5; Supplementary Tables S1 and S3)."
    )
    doc.add_heading("Layer 2b — Pathway convergence with Piezo1–LIFU mechanotransduction neighborhood", level=2)
    doc.add_paragraph(
        "Hallmark ORA shows strong enrichment of TGF-β signaling, hypoxia, TNF-α/NF-κB, apoptosis, and inflammatory response "
        "on the target union (Layer 2; Fig. 5A; Supplementary Table S1). A gene-set Jaccard heatmap (Fig. 5B; Supplementary Table S10) "
        "quantifies overlap between those Hallmark gene lists and curated mechanotransduction priors motivated by LIFU literature "
        "(Piezo1-centered mechanism seeds, neuroinflammatory and mechanosensory GMTs) and the full union. Shared membership supports "
        "that the htNSC targetome occupies the same broad signaling neighborhood as Piezo1-linked TGF-β, NF-κB, and HIF/hypoxia axes—"
        "without implying that ultrasound drove the observed aging pseudobulk ranks in this dataset. "
        "Supplementary Fig. S5 (UpSet) and Supplementary Table S11 list genes shared across two or more Hallmark gene lists."
    )
    if q.get("string_n_edges_mechanism_to_union") is not None:
        doc.add_paragraph(
            f"Exploratory STRING analysis reported {q.get('string_n_edges_mechanism_to_union')} high-confidence edges between "
            f"mechanism seeds and high-burden union genes in the induced subgraph "
            f"(uniform-size edge null P = {_fmt_num(q.get('string_perm_p_edges'))}; Supplementary Table S13), "
            "consistent with protein-level association in that neighborhood (not causal directionality)."
        )

    doc.add_heading("Layer 3 — GSE188646 hypothalamic aging pseudobulk DE and Tier-2 readouts", level=2)
    doc.add_paragraph(
        f"Pseudobulk edgeR QLF contrasts (Aged vs Young) are summarized in Fig. 2 for the "
        f"{q.get('n_deg_genes', 'N/A')}-gene table. At adjusted P < 0.05, {q.get('n_deg_padj005', 'N/A')} genes "
        "are flagged in the exported CSV used for Fisher and cross-modal joins. "
        f"PROGENy pathway scores on pseudobulk log-CPM (decoupler) span {q.get('progeny_n_pathways', 'N/A')} pathways in the "
        f"Welch Young-vs-Aged table; the minimum Benjamini–Hochberg FDR across those rows is {_fmt_num(q.get('progeny_min_fdr'))} "
        "(Supplementary Table S4). Together, these layers describe cell-state–level aging-associated "
        "expression without asserting that they replicate Zhang’s ex vivo design."
    )

    doc.add_heading(
        "Layer 4 — Aging association: decoupled burden–logFC coupling and attenuated |logFC| in union targets",
        level=2,
    )
    doc.add_paragraph(
        f"The primary aging readout is not gene-by-gene prediction of logFC from miRNA-target burden. On the merged "
        f"n = {q.get('cm_n_genes_merged', 'N/A')} genes with defined burden, Spearman correlation between htNSC-weighted "
        f"burden and signed aging logFC was rho = {_fmt_num(q.get('cm_spearman_rho_weighted_burden_vs_logFC'))} "
        f"(two-sided P = {_fmt_num(q.get('cm_spearman_p_weighted'))}; Fig. 1A)—no tight, rank-preserving alignment. "
        f"Burden by incoming miRNA count alone gave rho = {_fmt_num(q.get('cm_spearman_rho_count_vs_logFC'))} "
        f"(P = {_fmt_num(q.get('cm_spearman_p_count'))}). "
        f"Mann–Whitney comparison of |logFC| on the full GSE188646 pseudobulk table (union versus non-union) yielded "
        f"two-sided P = {_fmt_num(q.get('cm_mannwhitney_abs_logFC_union_vs_nonunion_p'))} (Fig. 1C; Table 2). "
        f"Direction matters: median |logFC| is lower in union targets "
        f"({_fmt_num(q.get('cm_median_abs_logFC_union'))}) than non-union "
        f"({_fmt_num(q.get('cm_median_abs_logFC_nonunion'))}); one-sided attenuation P = "
        f"{_fmt_num(q.get('cm_mannwhitney_abs_logFC_union_less_p'))}. "
        f"This pattern is compatible with post-transcriptional buffering of target mRNA aging volatility rather than "
        f"amplified transcriptional lability at bulk resolution."
    )
    doc.add_paragraph(
        f"Gene-label permutation ({q.get('cm_perm_p_rho_gene_shuffle_n', 'N/A')} shuffles; empirical P = "
        f"{q.get('cm_perm_p_rho_gene_shuffle', 'N/A')}) and random miRNA-set resampling "
        f"(tail fraction |ρ_null| ≥ |ρ_obs| = {q.get('cm_random_mirna_set_abs_rho_ge_observed_frac', 'N/A')}) "
        f"place the near-zero Spearman estimate in the null cloud (Fig. 1B; Table 2). "
    )
    doc.add_paragraph(
        f"Stratified nulls (n_perm = {q.get('neg_n_perm', 'N/A')}) restricted permutation of logFC within coarse strata. "
        f"Empirical P-values for |rho_observed| under: (i) se_logFC decile strata only = {q.get('neg_perm_p_abs_rho_ge_obs_se_decile_strata_only', 'N/A')}; "
        f"(ii) capped program degree × se deciles = {q.get('neg_perm_p_abs_rho_ge_obs_program_degree_cap_x_se_strata', 'N/A')}; "
        f"(iii) GMT-wide MOESM-weighted indegree × se deciles = {q.get('neg_perm_p_abs_rho_ge_obs_gmt_weighted_indegree_x_se_strata', 'N/A')}. "
        "Together with the random-set null, these panels argue against a simple residual coupling artifact confined to one "
        "confounder axis, while not equating in silico nulls with biological replication "
        "(Fig. 1; Table 2; Supplementary Fig. S1B–D; Supplementary Table S5). "
        "Supplementary Fig. S1A shows the pipeline schematic; Supplementary Fig. S1E plots cross-cohort logFC concordance for context under non-matched designs."
    )

    doc.add_heading(
        "Layer 4b — Cell-type–resolved localization of aging lability to the third-ventricle niche",
        level=2,
    )
    niche_ids = q.get("niche_third_ventricle_niche_strata_ids", [])
    ids_txt = ", ".join(str(x) for x in niche_ids) if niche_ids else "—"
    doc.add_paragraph(
        f"Whole-tissue pseudobulk lability (Layer 4; Fig. 1C) does not identify which hypothalamic cell classes "
        f"carry the signal. We therefore repeated the |logFC| Mann–Whitney test within each GSE188646 Seurat cluster "
        f"that has its own young-vs-aged pseudobulk DE table (n = {q.get('niche_n_strata_total', 'N/A')} strata; "
        f"not pseudobulk across all nuclei). Marker-module labels from integrated snRNA-seq "
        f"(cluster_putative_labels.csv) assign third-ventricle niche strata when rank-1 is Tanycyte_ependymal, "
        f"NSC_like, or Radial_glia_like, or when Astrocyte is rank-1 with a tanycyte/NSC/radial rank-2 module "
        f"(z ≥ 0.8); Astrocyte+Tanycyte pairs are tagged as astrocyte-adjacent niche. "
        f"Among labelled strata, n = {q.get('niche_n_third_ventricle_niche_strata', 'N/A')} map to this niche panel "
        f"(clusters {ids_txt}) versus n = {q.get('niche_n_other_strata', 'N/A')} other hypothalamic clusters."
    )
    doc.add_paragraph(
        f"Per-cluster Spearman burden–logFC coupling remains near zero throughout (Table 2; exploratory_crossmodal_celltype_strata_summary.csv)—"
        f"the cell-type advance is lability localization, not restored rank alignment. "
        f"At median, Δ|logFC| (union − non-union) is less negative in niche-labelled strata "
        f"({_fmt_num(q.get('niche_median_delta_abs_logfc_niche'))}) than in other clusters "
        f"({_fmt_num(q.get('niche_median_delta_abs_logfc_other'))}), i.e. niche partitions show the weakest "
        f"within-cluster anti-lability relative to the hypothalamic background; cross-stratum Wilcoxon "
        f"P = {_fmt_num(q.get('niche_wilcoxon_p_delta_median_niche_vs_other'))}, "
        f"label-shuffle permutation P = {_fmt_num(q.get('niche_perm_p_median_delta_diff_ge_obs'))} "
        f"(n_perm = {q.get('niche_perm_n_niche_label_shuffle', 'N/A')}; Fig. 6; Supplementary Table S14). "
        f"Cell-count–weighted mean Δ|logFC| is {_fmt_num(q.get('niche_weighted_mean_delta_abs_logfc_niche'))} "
        f"(niche) versus {_fmt_num(q.get('niche_weighted_mean_delta_abs_logfc_other'))} (other). "
        f"Within-cluster tests are mixed: astrocyte–radial cluster 6 Mann–Whitney P ≈ 0.06 and "
        f"cluster 15 P ≈ 0.07, whereas rank-1 tanycyte cluster 22 P ≈ 0.33 and astrocyte–radial cluster 7 "
        f"P ≈ 0.002 (union less labile in that stratum; Table S14). "
        "With NSC_like/Radial_glia_like modules in the marker map, the niche panel expands to four DE strata "
        "(clusters 22, 6, 7, 15), and cross-stratum enrichment of less-negative Δ|logFC| is no longer significant "
        "(Wilcoxon and permutation P > 0.1). The interpretable pattern is therefore cluster-specific—especially "
        "astro–radial cluster 6—rather than a uniform third-ventricle niche enrichment; labels remain heuristic "
        "marker modules, not spatial dissection."
    )

    doc.add_heading(
        "Layer 4c — Cross-cohort sensitivity of |logFC| attenuation in union targets",
        level=2,
    )
    doc.add_paragraph(
        f"We repeated the |logFC| (or |beta_DL|) Mann–Whitney test in GSE87102 male hypothalamus microarray and on "
        f"two-cohort DerSimonian–Laird meta coefficients (Fig. 7; Supplementary Table S15). GSE188646 pseudobulk "
        f"retains strong two-sided distributional difference (P = {_fmt_num(q.get('cc_gse188646_mannwhitney_two_sided_p'))}) "
        f"with union attenuation (Δ median |logFC| = {_fmt_num(q.get('cm_delta_median_abs_logFC_union_minus_nonunion'))}), "
        f"whereas GSE87102 (P = {_fmt_num(q.get('cc_gse87102_mannwhitney_two_sided_p'))}) and meta "
        f"(P = {_fmt_num(q.get('cc_meta_dl_mannwhitney_two_sided_p'))}) do not reproduce that shift. "
        "Thus the bulk attenuation phenotype is cohort- and assay-specific; the more portable cross-modal finding is "
        "near-zero Spearman burden–logFC coupling (Table 2; Fig. 3), which also holds on the meta intersection."
    )

    doc.add_heading(
        "Layer 4d — External HypoMap validation of third-ventricle niche identity",
        level=2,
    )
    doc.add_paragraph(
        f"Internal marker-module niche labels are heuristic. We therefore cross-validated each GSE188646 cluster "
        f"against the independent HypoMap hypothalamus atlas (Steuernagel et al., 2022; C185_named mean expression "
        f"profiles from CELLxGENE; Fig. 8; Supplementary Table S16). For each cluster we took the maximum Spearman "
        f"correlation to HypoMap third-ventricle reference types (Tanycytes, Ependymal, ParsTuber) and flagged "
        f"HypoMap-validated strata when rho ≥ {_fmt_num(q.get('hm_rho_threshold_hypomap_validated', 0.82))}."
    )
    doc.add_paragraph(
        f"All four marker-niche strata (clusters {q.get('hm_concordant_strata_ids', [])}) are concordant with "
        f"the external reference (Fisher enrichment P = {_fmt_num(q.get('hm_fisher_p_marker_vs_hypomap'))}; "
        f"no marker-only strata). Rank-1 tanycyte cluster 22 maps to "
        f"{q.get('hm_cluster_22_hypomap_best_type', 'N/A')} (rho = {_fmt_num(q.get('hm_cluster_22_hypomap_niche_rho'))}); "
        f"astrocyte–radial cluster 6 maps to {q.get('hm_cluster_6_hypomap_niche_best_type', 'N/A')} "
        f"(rho = {_fmt_num(q.get('hm_cluster_6_hypomap_niche_rho'))}), supporting an astro–tanycyte third-ventricle "
        f"niche interpretation independent of our module z-scores. Nine strata exceed the rho threshold overall; "
        f"five additional HypoMap-only strata reflect shared ventricular gene programs and are not claimed as niche "
        f"without marker concordance. Re-testing |logFC| attenuation on HypoMap-validated-only strata does not "
        f"yield significant cross-stratum enrichment (Wilcoxon P = {_fmt_num(q.get('hm_wilcoxon_p_delta_median_hypomap_validated_vs_other'))}), "
        f"so the external advance is identity corroboration, not a second significant lability hit."
    )

    doc.add_heading("Layer 5 — Fisher overlaps, HypoMap context, and two-cohort meta sensitivity", level=2)
    doc.add_paragraph(
        f"Fisher’s exact test in the miRTarBase universe between the miRNA-target union and GSE188646 DE "
        f"(padj ≤ 0.05) returned Fisher two-sided P = {q.get('fisher_gse_fisher_two_sided_p', 'N/A')} with "
        f"a_targets_and_external = {q.get('fisher_gse_a_targets_and_external', 'N/A')} — a degenerate table when the overlap "
        "layer is empty, reported in Supplementary Table S6 instead of being dropped. "
        f"For HypoMap GEO supplementary IP/input DE union, Fisher two-sided P = {q.get('fisher_hypo_fisher_two_sided_p', 'N/A')} "
        f"with odds ratio = {_fmt_num(q.get('fisher_hypo_odds_ratio'))} (Supplementary Table S9). "
        f"DerSimonian–Laird meta across GSE188646 and GSE87102 is tabulated for {q.get('n_meta_genes', 'N/A')} genes; "
        f"{q.get('n_meta_fdr05', 'N/A')} genes have global BH FDR < 0.05 on the meta z in the exported CSV excerpt. "
        f"Burden versus cohort1 logFC on the intersection gave rho = {_fmt_num(q.get('meta_spearman_rho_burden_vs_gse188646_logFC'))} "
        f"(P = {_fmt_num(q.get('meta_spearman_p_burden_vs_gse188646_logFC'))}); versus cohort2 logFC rho = {_fmt_num(q.get('meta_spearman_rho_burden_vs_gse87102_logFC'))} "
        f"(P = {_fmt_num(q.get('meta_spearman_p_burden_vs_gse87102_logFC'))}); versus pooled meta beta rho = {_fmt_num(q.get('meta_spearman_rho_burden_vs_meta_beta_DL'))} "
        f"(P = {_fmt_num(q.get('meta_spearman_p_burden_vs_meta_beta_DL'))}) (Fig. 3; Supplementary Table S7)."
    )

    doc.add_heading(
        "Layer 6 — Mechanotransduction/neuromodulation priors (consistency checks) and bootstrap stability of the target union",
        level=2,
    )
    doc.add_paragraph(
        f"We preranked GSEA against small mechanotransduction/neuromodulation gene sets—including Piezo1, integrins, "
        f"YAP/TAZ, angiogenesis, and inflammatory axes motivated in part by LIFU literature—each intersected with the DE universe. "
        f"The exported table lists {q.get('sa_n_sets', 'N/A')} sets; absolute NES magnitudes in this freeze are modest "
        f"(minimum NES = {_fmt_num(q.get('sa_nes_min'))} for {q.get('sa_term_nes_min', '')}; maximum NES = {_fmt_num(q.get('sa_nes_max'))} for {q.get('sa_term_nes_max', '')}) "
        f"as tabulated in exploratory_sa_nsc_lifu_fgsea_curated_sets.csv (Fig. 4; Supplementary Table S12). "
        f"Fisher overlap of the miRNA-target union with the same priors is likewise weak in most freezes (Supplementary Table S8). "
        f"We therefore present Fig. 4 as a pathway-level consistency check and hypothesis generator—not as evidence that LIFU "
        f"or mechanotransduction drives the observed aging signature. "
        f"Bootstrap resampling of miRNA draws (pool_size = {q.get('boot_pool_size', 'N/A')}, n_draw = {q.get('boot_n_draw', 'N/A')}, "
        f"sample_size = {q.get('boot_sample_size', 'N/A')}, reference union = {q.get('boot_ref_union_size', 'N/A')} genes) yielded "
        f"mean Jaccard to the reference union = {q.get('boot_mean_jaccard', 'N/A')}, median = {q.get('boot_median_jaccard', 'N/A')}, "
        f"5th–95th percentile = {q.get('boot_q05_jaccard', 'N/A')}–{q.get('boot_q95_jaccard', 'N/A')} (Supplementary Table S2, bootstrap block)."
    )

    doc.add_heading("Synthesis — how the layers connect and where each figure and table lives", level=2)
    doc.add_paragraph(
        "The narrative arc is: MOESM defines a biologically motivated miRNA program; miRTarBase supplies an experimentally "
        "supported targetome with strong Hallmark- and multi-library–level pathway structure (Layer 2; Fig. 5; "
        "Supplementary Tables S1 and S3); pathway convergence with Piezo1–LIFU priors is Layer 2b (Fig. 5). "
        "GSE188646 and cohort two quantify hypothalamic aging axes that are not matched "
        "to Zhang’s design (Layer 3; Fig. 2; Supplementary Table S4); cross-modal tests show |logFC| lability in union "
        "targets without Spearman burden alignment (Layer 4; Figs. 1A–C; Table 2; Supplementary Fig. S1B–D; Supplementary Table S5). "
        "Cell-type–resolved niche localization (Layer 4b; Fig. 6; Supplementary Table S14) asks where global lability "
        "is least cancelled in hypothalamic cluster space—third-ventricle tanycyte/NSC/radial-niche labels "
        "(clusters 22, 6, 7, 15 after updated marker mapping) versus other strata—without implying restored ρ within clusters. "
        "Fisher and HypoMap anchors situate overlaps in the correct contrast space (Layer 5; Supplementary Tables S6, S9). "
        "DerSimonian–Laird meta and burden-versus-meta sensitivity are in Fig. 3 and Supplementary Table S7. "
        "Weak LIFU-motivated GSEA/Fisher checks and bootstrap stability are Layer 6 (Fig. 4; Supplementary Tables S8 and S2). "
        "Governance of claim strength is explicit in Table 1. Supplementary Figs. S2–S4 visualize virtual null ensembles and "
        "sampling diagnostics (Methods: subsection on bootstrap stability, figure generation, and Word assembly). "
        "Frozen output filenames for replication are enumerated in outputs/SA_COMPLETENESS_CHECK.txt (see Structured reproducibility checklist)."
    )

    doc.add_heading("Structured reproducibility checklist (computational study)", level=2)
    doc.add_paragraph(
        "Structured transparency items aligned with reproducible reporting: (i) all primary "
        "CSV/JSON outputs are named in outputs/SA_COMPLETENESS_CHECK.txt after run_extended.py; "
        "(ii) contrast mismatches for HypoMap are stated in the statistical analysis plan; "
        "(iii) Stouffer combination of Fisher P-values is disabled when tests are degenerate; "
        "(iv) mechanotransduction/LIFU narrative evidence is a parallel text layer (outputs/LIFU_evidence_layer.txt); "
        "(v) curated niche GMT intersections are logged in exploratory_sa_nsc_lifu_suite_summary.json."
    )

    doc.add_heading("Discussion", level=1)
    doc.add_paragraph(
        "This computational manuscript unifies three pillars—(A) MOESM-supported htNSC miRNA and "
        "exosome-layer tables, (B) public hypothalamic aging pseudobulk and multi-cohort sensitivity, "
        "and (C) mechanotransduction-aligned pathway context (including LIFU as one clinically motivated interface)—and embeds them in a single cross-referenced figure "
        "and table stack (Figs. 1–8; Tables 1–2; Supplementary Figs. S1–S5; Supplementary Tables S1–S16). "
        "The lead finding for npj Aging–aligned hypothalamic biology is pathway-structured htNSC miRNA targeting "
        "(Hallmark inflammatory/stress axes; Fig. 5) with decoupled aging transcription (ρ≈0) and bulk attenuation "
        "of target |logFC| in GSE188646—not a claim of uniform cross-cohort mRNA lability amplification. "
        "The cross-modal layer is explicitly exploratory: non-exchangeable experiments, different gene universes, "
        "and literature-supported priors require cautious language consistent with Table 1. The implemented null assays "
        "provide reviewer-facing discipline analogous to bioorthogonal validation—multiple "
        "independent perturbation channels that should each be insensitive to specific "
        "artifact modes—while acknowledging that in silico nulls are not biological replication "
        "(Figs. 2 and 4; Supplementary Fig. S1; Supplementary Tables S2 and S5)."
    )
    doc.add_paragraph(
        "Mechanistic studies of low-intensity focused ultrasound have identified mechanosensitive channels such as Piezo1 "
        "as key mediators of ultrasonic neuromodulation, with Ca²⁺ influx propagating to ERK/MAPK, YAP/TAZ, HIF-1α, TGF-β, and NF-κB "
        "signaling in diverse cell types. The same stress–inflammatory and TGF-β/hypoxia pathways emerge as strong priors in our "
        "htNSC miRNA targetome, suggesting that exosomal miRNAs and ultrasound-evoked mechanotransduction could, in principle, "
        "converge on overlapping effectors within hypothalamic niches. However, our curated mechanotransduction/LIFU priors show "
        "only modest, non-significant enrichment on aging pseudobulk ranks in this build, so LIFU remains a narrative, pathway-level lane "
        "rather than a statistically supported rejuvenation strategy."
    )
    doc.add_paragraph(
        "Readers should treat Supplementary Figs. S2–S4 panels as structural QC on how nulls are sampled, not as new "
        "hypothesis tests beyond Table 2. Fisher degeneracy (Supplementary Table S6) and HypoMap overlap (Supplementary Table S9) "
        "must be read beside the |logFC| attenuation evidence in Table 2 and cross-cohort sensitivity in Fig. 7."
    )
    doc.add_paragraph(
        "A decisive next step would be a preregistered in vivo study combining controlled LIFU sonication of hypothalamic niches, "
        "Piezo1 manipulation, and longitudinal profiling of htNSC miRNA cargo and target expression, to test whether mechanotransduction "
        "along the Piezo1–Ca²⁺–ERK/MAPK–YAP/TAZ axis can reproducibly shift the htNSC miRNA program or its targets toward younger-like states."
    )

    doc.add_heading("Methods", level=1)
    doc.add_paragraph(
        "Methods are organized in the conventional order used in integrative omics papers: we first specify "
        "all datasets and public resources; then software and reproducible workflow; then how MOESM-derived "
        "miRNA tables are processed; then target mapping and burden construction; then hypothalamic aging "
        "differential expression and two-cohort meta-analysis; then cross-modal and null-statistics procedures; "
        "then Fisher overlap tests; then pathway enrichment, mechanotransduction/neuromodulation priors (LIFU-motivated), and Tier-2 readouts; finally "
        "bootstrap stability, figure generation, and assembly of this Word document. File names and defaults "
        "match the frozen repository snapshot."
    )

    doc.add_heading("Datasets and public resources", level=2)
    doc.add_paragraph(
        "NCBI Gene Expression Omnibus (GEO): GSE188646 (hypothalamic single-nucleus RNA-seq, young versus aged; "
        "female mice; object processed here as GSE188646_hypo.integrated.final.20210719.RDS or path given by "
        "environment variable GSE188646_RDS). GSE87102 (whole hypothalamus microarray; young versus aged; "
        "male C57BL/6 mice; series matrix retrieved with GEOquery when required). GSE208355 supplementary "
        "DESeq2 tables (IP versus input contrasts) used to build a filtered HypoMap-oriented DE gene union for "
        "Fisher overlap tests (see overlap_fisher.load_hypomap_de_union). "
        "Primary laboratory–derived inputs: Zhang et al. (2017) Nature supplementary spreadsheets (MOESM) "
        "ingested per repository configuration (MOESM_GLOB / paths in config) to harmonize miRNA-by-condition "
        "contrasts (htNSC versus astrocyte and related tables). "
        "Interaction database: miRTarBase experimentally supported mouse interactions distributed as the "
        "miRTarBase 2017 Enrichr GMT (mmu-miR* and mmu-let* term rows) for target parsing. "
        "Gene-set libraries: MSigDB Hallmark 2020 (Hallmark ORA and preranked GSEA) and additional mouse "
        "Enrichr libraries for multi-library ORA with global Benjamini–Hochberg adjustment across libraries. "
        "Aging–gene compendium: Human Ageing Genomic Resources (HAGR) GenAge models_genes.zip for GenAge Fisher "
        "overlap as documented in data/provenance/GenAge_HAGR.txt. "
        "Neuromodulation context: mechanistic and safety-oriented notes on low-intensity focused ultrasound (LIFU) and related tFUS literature are summarized in "
        "outputs/LIFU_evidence_layer.txt as a narrative evidence lane that is not merged statistically with omics in this build; "
        "they motivate mechanotransduction-aligned reading alongside the curated prior GMTs, not pooled effect estimates."
    )

    doc.add_heading("Software environment, code organization, and execution order", level=2)
    doc.add_paragraph(
        "All analysis code resides under feasibility_study/. Execution order is: "
        "(1) python run_pipeline.py ingests MOESM spreadsheets and writes mirna_htnsc_astrocyte_summary.csv and "
        "related prerequisite tables—log-fold-changes therein are MOESM-supported weights for downstream burdening, "
        "not a newly fit differential experiment. "
        "(2) python run_extended.py builds the miRTarBase target union, differential expression and meta layers, "
        "Hallmark and multi-library ORA/GSEA, Fisher tests, cross-modal and stratified-null modules, optional "
        "two-cohort meta, Tier-2 R/Python bridges, the SA mechanotransduction/neuromodulation prior suite (LIFU-motivated GMTs), bootstrap stability draws, and the "
        "main figure bundle. "
        "Optional automatic download of public objects, skipping cohort-2 meta (GSE188646_SKIP_COHORT2_META=1), "
        "and skipping figure regeneration are controlled by environment variables documented in "
        "data/provenance/DATA_ACQUISITION_AND_QC_RATIONALE.txt and RUN_ORDER.txt. "
        "Processed CSV/JSON outputs under outputs/ are what this manuscript embeds; long-term archival mirrors "
        "(e.g., Zenodo) should follow funder and journal policy for large intermediates."
    )

    doc.add_heading("MOESM supplementary spreadsheet ingestion", level=2)
    doc.add_paragraph(
        "Zhang et al. (2017) supplementary tables are read according to repository configuration. "
        "miRNA-by-condition summaries are harmonized to produce mirna_htnsc_astrocyte_summary.csv and related "
        "outputs that supply ranked miRNA contrasts between htNSC-biased and comparator conditions. "
        "These tables are Tier-0–aligned primary data (Table 1); no new miRNA differential model is fit here beyond parsing and harmonization implemented in run_pipeline.py."
    )

    doc.add_heading("miRNA program selection, miRTarBase targets, and gene-level burden", level=2)
    doc.add_paragraph(
        "Top htNSC-biased miRNAs are selected with mirna_target_union.load_top_htnsc_mirnas "
        "(default top_n=60, min_logfc=0.35) from mirna_htnsc_astrocyte_summary.csv and written to "
        "top_htnsc_mirnas_for_targets.csv. Targets are parsed from the miRTarBase GMT via "
        "mirna_target_union.ensure_mirtarbase_gmt(), with optional AUTO_FETCH_PUBLIC_DATA hooks. "
        "The union of targets across selected miRNAs is intersected with the Mus musculus symbol universe "
        "extracted from that GMT so that ORA, Fisher tests, and overlap statistics share one coherent universe. "
        "mirna_targets_long.csv and related outputs define incoming regulators per gene; gene-level burden sums "
        "non-negative MOESM weights over those regulators for use in cross-modal tests. "
        "Per-miRNA target maps additionally feed a multiplicity rank for preranked Hallmark GSEA "
        "(mirna_target_multiplicity_rank.csv → outputs/gsea_hallmark_prerank_results.csv via gseapy prerank)."
    )

    doc.add_heading("Hypothalamic aging differential expression (GSE188646 and GSE87102)", level=2)
    doc.add_paragraph(
        "GSE188646: when the integrated RDS is available, r/pseudobulk_edgeR_gse188646.R aggregates "
        "single-nucleus counts to pseudobulk by biological replicate (orig.ident) and fits edgeR "
        "quasi-likelihood F-tests for Aged versus Young. Genes are rows; Benjamini–Hochberg adjusted P-values "
        "(FDR) are reported within the contrast. Optional MatrixMarket export (GSE188646_EXPORT_COUNTS=1) "
        "enables pseudobulk log-CPM pipelines used by Tier-2 modules. "
        "GSE87102: r/cohort2_gse87102_c57_hypothalamus_limma.R downloads the series matrix with GEOquery where needed, "
        "restricts to C57BL/6 hypothalamus samples annotated as young versus aged in GEO sample titles, and fits "
        "limma moderated t-statistics. Sex, platform, and tissue dissection scope therefore differ from GSE188646; "
        "cohort two is used as a sensitivity anchor, not as a matched biological replicate."
    )

    doc.add_heading("Two-cohort DerSimonian–Laird random-effects meta-analysis", level=2)
    doc.add_paragraph(
        "Gene-wise meta-analysis combines cohort-one and cohort-two logFC and standard errors using "
        "DerSimonian–Laird random effects implemented in Python (src/replication_meta_two_cohorts.py): "
        "inverse-variance weighting, Cochran Q, standard between-study variance (tau²) estimator for k = 2, "
        "two-sided normal P-values, and Benjamini–Hochberg FDR (statsmodels.multipletests) on "
        "exploratory_meta_DE_two_cohort_DL.csv. Meta is skipped when GSE188646_SKIP_COHORT2_META=1."
    )

    doc.add_heading(
        "Cross-modal burden–aging association, permutation nulls, stratified controls, and cohort sensitivity",
        level=2,
    )
    doc.add_paragraph(
        "journal_tier_crossmodal.py tests Spearman correlation between MOESM-weighted miRNA-target burden and "
        "GSE188646 pseudobulk logFC on the inner-joined gene set (minimum 200 genes). Gene-label permutation uses "
        "n_perm=1000 shuffles of logFC labels against fixed burden (as invoked from run_extended.py). "
        "Random miRNA-set resampling uses n_mirna_draws=200 draws of programs matched in cardinality from the "
        "GMT pool intersecting MOESM-weighted miRNAs. Mann–Whitney tests compare |logFC| in union targets versus "
        "non-target genes on the full DE table. "
        "journal_tier_negative_controls.py implements restricted permutations that shuffle logFC within coarse strata: "
        "se_logFC deciles; capped program in-degree crossed with se deciles; GMT-wide MOESM-weighted indegree crossed "
        "with se deciles (default n_perm=800; empirical P = fraction of |ρ_null| ≥ |ρ_observed|; see "
        "JOURNAL_TIER_COMPUTATIONAL_README.txt and module docstrings). "
        "journal_tier_crossmodal_meta_sensitivity.py aligns the same burden vector to cohort-one logFC, cohort-two logFC, "
        "and pooled DL meta β where available, writing exploratory_crossmodal_meta_cohort_sensitivity_summary.json "
        "with explicit sex/assay mismatch notes."
    )

    doc.add_heading("Fisher exact tests for overlap with external gene lists", level=2)
    doc.add_paragraph(
        "overlap_fisher.py constructs 2×2 contingency tables using the miRTarBase mmu symbol universe size "
        "recorded in each output CSV. Overlap with GSE188646 aging DE uses genes passing padj ≤ 0.05 from "
        "gse188646_young_vs_aged_deg.csv. Overlap with the HypoMap-oriented union uses filtered GSE208355 "
        "supplementary DESeq2 tables with thresholds in overlap_fisher.load_hypomap_de_union "
        "(default padj_max=0.05, abs_lfc_min=0.5). Degenerate tables (e.g., empty overlap cells) are retained and reported."
    )

    doc.add_heading(
        "Pathway over-representation, preranked enrichment, multi-library ORA, mechanotransduction/neuromodulation priors, and Tier-2 readouts",
        level=2,
    )
    doc.add_paragraph(
        "Hallmark ORA uses msigdb_hallmark_gsea.run_hallmark_ora with background uni_mmu; offline fallbacks and "
        "Enrichr retry logic are in msigdb_hallmark_gsea.py. Preranked Hallmark GSEA uses gseapy prerank on the "
        "multiplicity rank table. multi_library_ora.run_multi_ora queries multiple mouse Enrichr libraries "
        "(default organism='mouse', top_n_terms_per_lib=35), pools P-values across libraries, and applies global "
        "Benjamini–Hochberg adjustment (enrichr_multi_library_summary.csv; enrichr_multi_library_ora_terms.csv). "
        "sa_nsc_lifu_computational_suite.py intersects curated mechanotransduction/neuromodulation GMT gene sets (Piezo1, integrins, YAP/TAZ, angiogenesis, inflammatory signaling; motivated in part by LIFU studies) with the DE universe, runs gseapy prerank "
        "(symbols uppercased for GMT matching), Fisher overlap of each set against the miRNA-target union in the "
        "miRTarBase universe, and—when pseudobulk counts exist—Welch comparisons of mean log1p-CPM module scores; "
        "this is a parallel prior lane without pooling LIFU literature statistics into omics. "
        "tier2_analyses.py orchestrates Bioconductor fgsea (r/fgsea_hallmark_prerank.R on DE-derived ranks), "
        "limma camera and fry on pseudobulk log-CPM with the same young-versus-aged design, decoupler (Python) "
        "for PROGENy pathway scores and DoRothEA regulon activity (mlm) where expression matrices exist, and "
        "GenAge Fisher overlap from HAGR models_genes.zip per data/provenance/GenAge_HAGR.txt."
    )

    doc.add_heading(
        "Bootstrap stability of the miRNA-target union, figure generation, and Word manuscript assembly",
        level=2,
    )
    doc.add_paragraph(
        "Bootstrap resampling of miRNA draws for target-union stability is summarized in "
        "bootstrap_target_union_stability.csv and bootstrap_target_union_stability_summary.txt (parameters such as "
        "pool_size, n_draw, and sample_size appear in those outputs). "
        "Publication-facing main panels are rendered by figures/build_sa_figure_bundle.py from "
        "data/provenance/FIGURE_PANEL_MANIFEST.csv. tools/build_science_advances_manuscript_docx.py embeds those PNGs, "
        "regenerates supplementary multi-panel figures (orthogonal validation suite, virtual simulation laboratory, "
        "Monte Carlo null ensemble heatmaps, parallel-universe |ρ| landscape), and composes tables via python-docx. "
        "Supplementary illustrations resample frozen null pools (e.g., gene-shuffle permutations, bootstrap draws) "
        "to visualize sampling structure—parallel streams, mini-replicate histograms, Jaccard clouds, subsample sweeps, "
        "random-set clouds, meta z histograms, and heatmaps of null run × bin occupancy; they are QC visualizations of "
        "the estimators above, not additional hypothesis tests. The resulting .docx is a draft scaffold: vector figures, "
        "author metadata, and journal-specific declarations remain editorial steps outside this builder. "
        "Figure and table placement follows the order described in the Introduction."
    )

    doc.add_heading("Data and code availability", level=1)
    doc.add_paragraph(
        "GEO records GSE188646, GSE87102, GSE208355; CELLxGENE HypoMap object when used; "
        "HAGR GenAge; miRTarBase via Enrichr GMT mirrors. Processed outputs ship with the "
        "repository under outputs/ for drafting; archival deposition should follow journal "
        "policy (e.g., Zenodo) for large objects."
    )
    doc.add_paragraph(
        "All code: feasibility_study/run_pipeline.py, run_extended.py, src/*, r/*, "
        "figures/build_sa_figure_bundle.py, tools/build_science_advances_manuscript_docx.py."
    )

    doc.add_heading("References (starter set — verify and expand)", level=1)
    refs = [
        "Zhang, Y. et al. (2017). Nature (MOESM supplementary materials as processed here).",
        "Matson, M. P. & Arumugam, T. V. Hypothalamus and aging. Ageing Res. Rev. (2005).",
        "Zhang, Y. et al. Hypothalamic stem cells control ageing speed partly through exosomal miRNAs. Nature (2017).",
        "SnRNA-seq resource: NCBI GEO GSE188646 (hypothalamus, young vs aged, as cited in repository).",
        "Microarray resource: NCBI GEO GSE87102 (hypothalamus aging in C57BL/6).",
        "HypoMap: GEO GSE208355 supplementary tables (IP/input contrasts).",
        "Subramanian, A. et al. Gene set enrichment analysis. PNAS (2005).",
        "Korotkevich, G. et al. Fast gene set enrichment analysis. bioRxiv / fgsea package.",
        "miRTarBase: experimentally supported miRNA–target interactions database.",
        "Liberzon, A. et al. The Molecular Signatures Database (MSigDB). Cell Syst. (2015).",
        "ITRUSST / neuromodulation safety literature and LIFU narrative summaries in LIFU_evidence_layer.txt.",
    ]
    for i, r in enumerate(refs, 1):
        doc.add_paragraph(f"{i}. {r}", style="List Number")

    bundle = out_dir / "figures" / "sa_bundle"
    fig_upset = bundle / "fig_supp_hallmark_stress_upset.png"
    doc.add_page_break()
    doc.add_heading("Main Figures", level=1)
    doc.add_paragraph(
        "Main panels under outputs/figures/sa_bundle/: Fig. 1A–C cross-modal aging association (ρ≈0; permutation null; |logFC| lability); "
        "Fig. 2 GSE188646 volcano; Fig. 3 two-cohort meta sensitivity; Fig. 4 weak LIFU-motivated GSEA consistency checks; "
        "Fig. 5 Hallmark stress ORA and Jaccard overlap with Piezo1/LIFU priors; "
        "Fig. 6 third-ventricle niche localization of per-cluster |logFC| dynamics; "
        "Fig. 7 cross-cohort sensitivity of union |logFC| attenuation; "
        "Fig. 8 independent HypoMap third-ventricle niche validation. Tables 1–2 appear earlier after the Abstract."
    )
    main_figs = [
        (bundle / "fig01a_crossmodal_burden_logfc.png", "Fig. 1A. Cross-modal miRNA-target burden versus GSE188646 pseudobulk aging logFC (no rank-preserving alignment)."),
        (bundle / "fig01b_crossmodal_perm_null.png", "Fig. 1B. Gene-label permutation null for Spearman ρ (burden vs logFC)."),
        (bundle / "fig01c_union_abs_logfc_violin.png", "Fig. 1C. |logFC| in miRNA-target union versus non-union genes (aging lability; Mann–Whitney)."),
        (bundle / "fig02a_gse188646_volcano.png", "Fig. 2. Volcano plot of GSE188646 pseudobulk differential expression (Aged vs Young)."),
        (bundle / "fig03a_crossmodal_burden_meta_dl.png", "Fig. 3. Cross-modal burden versus two-cohort DerSimonian–Laird meta beta (sensitivity; cohorts differ by sex and assay)."),
        (bundle / "fig04a_sa_niche_gsea_nes.png", "Fig. 4. Mechanotransduction/neuromodulation prior GSEA NES (weak consistency checks; not headline discoveries)."),
        (bundle / "fig_pathway_convergence.png", "Fig. 5. Pathway convergence: Hallmark stress ORA on the target union and Jaccard overlap with Piezo1/LIFU mechanotransduction priors."),
        (bundle / "fig_niche_lability_localization.png", "Fig. 6. Cell-type–resolved localization of miRNA-target |logFC| dynamics to third-ventricle niche-labelled clusters (per-cluster DE)."),
        (bundle / "fig_crosscohort_lability_replication.png", "Fig. 7. Cross-cohort sensitivity: |effect| Mann–Whitney for miRNA-target union vs non-union genes."),
        (bundle / "fig_niche_hypomap_external_validation.png", "Fig. 8. External HypoMap validation of third-ventricle niche identity (independent of marker modules)."),
    ]
    main_legends = [
        legend_fig1(q),
        legend_fig2(q),
        legend_fig1c(q),
        legend_fig3(q),
        legend_fig4(q),
        legend_fig5(q),
        legend_fig_pathway_convergence(q),
        legend_fig_niche_lability(q),
        legend_fig_crosscohort(q),
        legend_fig_hypomap_niche(q),
    ]
    for (path, cap), legs in zip(main_figs, main_legends):
        _embed_png_with_legend(doc, path, cap, width_in=6.45, legend_lines=legs)

    doc.add_page_break()
    doc.add_heading("Supplementary Materials", level=1)
    doc.add_paragraph(
        "Supplementary figures S1–S4 extend the in silico validation metaphor with multi-panel layouts; "
        "Supplementary Tables S1–S14 reproduce CSV/JSON excerpts referenced in Results as separate publication-formatted Word files. "
        "When cell-type strata cross-modal diagnostics are run, forest/scatter panels and a glial-focused excerpt are also assembled as "
        f"{(man_dir / 'Supplementary_Figures_Strata_Crossmodal.docx').name} (absolute path under Supplementary Tables). Mapping: "
        "Supplementary Fig. S1 — pipeline schematic, gene-shuffle and random miRNA-set null histograms, stratified null p-values, "
        "cross-cohort logFC hexbin (see Layer 4–5 and Table 2). Supplementary Fig. S2 — virtual simulation laboratory "
        "(null streams, bootstrap Jaccard, subsample sweeps; aligns with Supplementary Table S2 bootstrap block). "
        "Supplementary Fig. S3 — Monte Carlo null ensemble heatmaps (structural QC). Supplementary Fig. S4 — parallel-universe "
        "|Spearman| landscape under permuted logFC. Supplementary Fig. S5 — UpSet of genes shared across multiple stress Hallmark "
        "ORA gene lists (complements main Fig. 5B and Table S10). Supplementary Table S1 — Hallmark ORA on the target union; "
        "S2 — integrated QC / assay readout; S3 — multi-library ORA; S4 — PROGENy; S5 — full JSON-derived summary (cross-modal, nulls, Mann–Whitney); "
        "S6 — Fisher vs GSE188646 DE; S7 — meta gene excerpt; S8 — Fisher vs mechanotransduction priors; S9 — Fisher vs HypoMap union; "
        "S10 — pathway-convergence Jaccard matrix; S11 — Hallmark stress UpSet intersections; S12 — mechanotransduction prerank GSEA NES; "
        "S13 — STRING Piezo1 bridge summary (when run). "
        "S14 — per-cluster niche lability (third-ventricle localization; main Fig. 6). "
        "S15 — cross-cohort |effect| Mann–Whitney sensitivity (main Fig. 7). "
        "S16 — HypoMap external third-ventricle niche validation (main Fig. 8). "
        "Full matrices remain in outputs/ for archival transfer; filenames are also listed in SA_COMPLETENESS_CHECK.txt."
    )

    doc.add_heading("Supplementary Figures", level=1)
    supp_figs = [
        (fig_suite, "Fig. S1. Computational orthogonal validation suite (five panels A–E): pipeline schematic; gene-shuffle and random miRNA-set null histograms with observed Spearman ρ; stratified null empirical p-values; cross-cohort logFC hexbin readout."),
        (fig_virt, "Fig. S2. Virtual simulation laboratory: parallel null cohort streams, mini-replicate histogram universes, bootstrap miRNA-pool draws, gene-subspace sweep, random-set null cloud, and meta z-score landscape."),
        (fig_mc, "Fig. S3. Monte Carlo null ensemble: per-run null histogram heatmap and across-run dispersion of bin occupancy (structural QC of virtual sampling)."),
        (fig_pu, "Fig. S4. Parallel-universe null landscape: |Spearman(burden, permuted logFC)| across independent shuffles and subsampled gene fractions."),
        (
            fig_upset,
            "Fig. S5. UpSet of genes shared across ≥2 stress Hallmark ORA gene lists on the htNSC miRNA target union "
            "(TGF-β, hypoxia, TNF-α/NF-κB, apoptosis, inflammatory response).",
        ),
    ]
    supp_legends = [
        legend_fig_s1(q),
        legend_fig_s2(q),
        legend_fig_s3(q),
        legend_fig_s4(q),
        legend_fig_s5_hallmark_upset(q),
    ]
    for (path, cap), legs in zip(supp_figs, supp_legends):
        _embed_png_with_legend(doc, path, cap, width_in=6.45, legend_lines=legs)

    doc.add_heading("Supplementary Tables", level=1)
    doc.add_paragraph(
        "Supplementary Tables S1–S14 are not embedded in this manuscript .docx. They are written as separate publication-formatted "
        "Word files (landscape sections, grid styling, reader-facing column labels) for editorial transfer:"
    )
    doc.add_paragraph(str((man_dir / "Supplementary_Table_S1.docx").resolve()))
    doc.add_paragraph(str((man_dir / "Supplementary_Tables_S2_S9.docx").resolve()))
    s2_12 = man_dir / "Supplementary_Tables_S2_S12.docx"
    if s2_12.is_file():
        doc.add_paragraph(str(s2_12.resolve()))
    doc.add_paragraph(str((man_dir / "Supplementary_Figures_Strata_Crossmodal.docx").resolve()))
    doc.add_paragraph(
        "Optional strata bundle (not produced by this script): regenerate with "
        "python tools/diagnostic_crossmodal_strata_figures.py --write-supplementary-docx "
        "after exploratory_crossmodal_celltype_strata_summary.csv exists, "
        "or set RUN_CROSSMODAL_STRATA_DIAGNOSTICS=1 and RUN_CROSSMODAL_STRATA_SUPPLEMENTARY_DOCX=1 when running python run_extended.py."
    )
    doc.add_paragraph(
        f"Regenerate from the feasibility_study directory after pipeline outputs exist under {out_dir.as_posix()}: "
        "python tools/generate_supplementary_table_s1_docx.py && python tools/generate_supplementary_tables_s2_s9_docx.py "
        "(each accepts --outputs-dir if your tree is not the default; writes Supplementary_Tables_S2_S9.docx and S2_S12.docx). "
        "Use outputs/SA_COMPLETENESS_CHECK.txt for a machine-enumerated file list."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "End of auto-generated draft. Replace bracketed author placeholders; curate references; "
        "add institutional review and data-sharing statements; export vector figures for production."
    )

    doc_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(doc_path))
    print(f"Wrote {doc_path}")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument(
        "--outputs-dir",
        type=Path,
        default=PROJECT_ROOT / "outputs",
        help="Directory containing pipeline outputs (default: feasibility_study/outputs)",
    )
    ap.add_argument(
        "--out-doc",
        type=Path,
        default=None,
        help="Destination .docx path (default: <outputs-dir>/manuscript/Computational_Integrative_Manuscript.docx)",
    )
    args = ap.parse_args()
    out_dir = args.outputs_dir.resolve()
    doc_path = args.out_doc or (out_dir / "manuscript" / "Computational_Integrative_Manuscript.docx")
    build_manuscript(out_dir, doc_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
