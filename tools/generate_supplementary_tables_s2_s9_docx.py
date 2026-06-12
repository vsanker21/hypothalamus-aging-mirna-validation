"""
Standalone publication-style Word document for Supplementary Tables S2–S14.

Table S1: use generate_supplementary_table_s1_docx.py.

Usage (from feasibility_study/):
  python tools/generate_supplementary_tables_s2_s9_docx.py

Output (same builder; extended table range):
  outputs/manuscript/Supplementary_Tables_S2_S9.docx
  outputs/manuscript/Supplementary_Tables_S2_S12.docx  (duplicate save for clarity)
"""
from __future__ import annotations

import argparse
import json
import numbers
import shutil
import sys
from pathlib import Path

import pandas as pd

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError as e:
    print("Install python-docx: pip install python-docx", file=sys.stderr)
    raise SystemExit(1) from e


PROJECT_ROOT = Path(__file__).resolve().parents[1]

_JSON_KEY_LABELS: dict[str, str] = {
    "n_genes_merged": "Genes in cross-modal merge (burden ∩ cohort-1 logFC)",
    "n_distinct_mirnas_in_program": "Distinct miRNAs in MOESM-weighted program",
    "spearman_rho_weighted_burden_vs_logFC": "Spearman ρ (weighted burden vs. GSE188646 logFC)",
    "spearman_p_weighted": "P value (weighted burden vs. logFC, Spearman)",
    "spearman_rho_count_vs_logFC": "Spearman ρ (count-based burden vs. logFC)",
    "spearman_p_count": "P value (count-based burden vs. logFC)",
    "perm_p_rho_gene_shuffle_n": "Gene-label permutation replicates (n_perm)",
    "perm_p_rho_gene_shuffle": "Empirical P (gene-label shuffle vs. Spearman ρ)",
    "random_mirna_set_draws": "Random miRNA-set draws (n)",
    "random_mirna_set_abs_rho_ge_observed_frac": "Tail fraction (|ρ_null| ≥ |ρ_observed|)",
    "mannwhitney_abs_logFC_union_vs_nonunion_p": "Mann–Whitney P (|logFC| union vs. non-union)",
    "n_genes": "Genes in stratified-null analysis",
    "spearman_rho_observed_weighted_burden_vs_logFC": "Observed Spearman ρ (stratified-null context)",
    "n_perm": "Stratified permutation replicates (n_perm)",
    "perm_p_abs_rho_ge_obs_se_decile_strata_only": "Stratified-null P (SE-decile strata only)",
    "perm_p_abs_rho_ge_obs_program_degree_cap_x_se_strata": "Stratified-null P (program degree × SE)",
    "perm_p_abs_rho_ge_obs_gmt_weighted_indegree_x_se_strata": "Stratified-null P (GMT indegree × SE)",
    "completed_utc": "Run timestamp (UTC)",
    "n_genes_intersection": "Genes at burden ∩ meta-table intersection",
    "spearman_rho_burden_vs_gse188646_logFC": "Spearman ρ (burden vs. GSE188646 logFC)",
    "spearman_p_burden_vs_gse188646_logFC": "P value (burden vs. GSE188646 logFC)",
    "spearman_rho_burden_vs_gse87102_logFC": "Spearman ρ (burden vs. GSE87102 logFC)",
    "spearman_p_burden_vs_gse87102_logFC": "P value (burden vs. GSE87102 logFC)",
    "spearman_rho_burden_vs_meta_beta_DL": "Spearman ρ (burden vs. DL meta β)",
    "spearman_p_burden_vs_meta_beta_DL": "P value (burden vs. DL meta β)",
}

_FISHER_ROW_LABELS: dict[str, str] = {
    "a_targets_and_external": "Overlap: targets and DE (a)",
    "b_targets_only": "Targets only (b)",
    "c_external_only": "DE / external only (c)",
    "d_neither": "Neither (d)",
    "universe_size": "Universe size (genes)",
    "odds_ratio": "Odds ratio",
    "fisher_two_sided_p": "Fisher two-sided P",
}


def _read_json(p: Path) -> dict | None:
    if not p.is_file():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None


def _safe_csv(p: Path) -> pd.DataFrame | None:
    if not p.is_file():
        return None
    try:
        return pd.read_csv(p, nrows=None)
    except Exception:
        return None


def _fmt_val(v: object) -> str:
    if v is None or (isinstance(v, float) and pd.isna(v)):
        return ""
    if isinstance(v, bool):
        return "Yes" if v else "No"
    if isinstance(v, numbers.Integral):
        return str(int(v))
    if isinstance(v, float):
        if abs(v) < 1e-6 or abs(v) >= 1e5:
            return f"{v:.4e}"
        return f"{v:.6g}".rstrip("0").rstrip(".")
    s = str(v)
    return s[:800] + ("…" if len(s) > 800 else "")


def _set_cell_text(cell, text: str, *, bold: bool = False, size_pt: float = 9, align=None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"


def _set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for el in tc_pr.findall(qn("w:tcW")):
        tc_pr.remove(el)
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _set_table_grid(table, col_widths_twips: list[int]) -> None:
    tbl = table._tbl
    for el in list(tbl.findall(qn("w:tblGrid"))):
        tbl.remove(el)
    tbl_grid = OxmlElement("w:tblGrid")
    for w in col_widths_twips:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tbl_grid.append(gc)
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl.insert(0, tbl_grid)
        return
    insert_at = list(tbl).index(tbl_pr) + 1
    tbl.insert(insert_at, tbl_grid)


def _set_table_full_width(table, width_twips: int) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    for el in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(el)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    for el in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(el)
    tl = OxmlElement("w:tblLayout")
    tl.set(qn("w:type"), "fixed")
    tbl_pr.append(tl)


def _repeat_header_row(table) -> None:
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    for el in tr_pr.findall(qn("w:tblHeader")):
        tr_pr.remove(el)
    tr_pr.append(OxmlElement("w:tblHeader"))


def _shade_header_row(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9E2F3")
        shd.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shd)


def _row_height_at_least(row, twips: int = 360) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for el in tr_pr.findall(qn("w:trHeight")):
        tr_pr.remove(el)
    rh = OxmlElement("w:trHeight")
    rh.set(qn("w:val"), str(twips))
    rh.set(qn("w:hRule"), "atLeast")
    tr_pr.append(rh)


def _apply_professional_table(
    table,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[int],
    col_align: list[str],
) -> None:
    assert len(headers) == len(col_widths) == len(col_align)
    total_w = sum(col_widths)
    table.style = "Table Grid"
    _set_table_full_width(table, total_w)
    _set_table_grid(table, col_widths)
    _repeat_header_row(table)
    am = {"l": WD_ALIGN_PARAGRAPH.LEFT, "c": WD_ALIGN_PARAGRAPH.CENTER, "r": WD_ALIGN_PARAGRAPH.RIGHT}
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        _set_cell_width(cell, col_widths[j])
        _set_cell_text(cell, h, bold=True, size_pt=9, align=am.get(col_align[j], WD_ALIGN_PARAGRAPH.CENTER))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _row_height_at_least(table.rows[0], 380)
    for ri, row_vals in enumerate(rows, start=1):
        for j, val in enumerate(row_vals):
            cell = table.rows[ri].cells[j]
            _set_cell_width(cell, col_widths[j])
            _set_cell_text(cell, val, size_pt=9, align=am.get(col_align[j], WD_ALIGN_PARAGRAPH.LEFT))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _row_height_at_least(table.rows[ri], 360)
    _shade_header_row(table)


def _json_key_label(k: str) -> str:
    return _JSON_KEY_LABELS.get(k, k.replace("_", " ").strip().title())


def _add_table_block(
    doc: Document,
    *,
    title: str,
    source_note: str,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[int],
    col_align: list[str],
    legend: str,
) -> None:
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    r = p.add_run(source_note)
    r.italic = True
    r.font.size = Pt(9)
    doc.add_paragraph()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _apply_professional_table(table, headers, rows, col_widths, col_align)
    doc.add_paragraph()
    lp = doc.add_paragraph()
    lr = lp.add_run("Table legend. ")
    lr.bold = True
    lr.font.size = Pt(10)
    lp.add_run(legend).font.size = Pt(9)


def _gather_s2_rows(out_dir: Path) -> tuple[list[list[str]], str]:
    cm0 = _read_json(out_dir / "exploratory_crossmodal_mirna_aging_summary.json")
    neg0 = _read_json(out_dir / "exploratory_negative_controls_summary.json")
    meta0 = _read_json(out_dir / "exploratory_crossmodal_meta_cohort_sensitivity_summary.json")
    rows: list[list[str]] = []
    if cm0:
        for k, v in cm0.items():
            if k == "caveat":
                continue
            rows.append(["Cross-modal summary", _json_key_label(k), _fmt_val(v)])
    if neg0:
        for k, v in neg0.items():
            if k in ("methodology_note", "caveat", "strata_notes"):
                continue
            rows.append(["Stratified nulls", _json_key_label(k), _fmt_val(v)])
    if meta0:
        for k, v in meta0.items():
            if k in ("methodology_note", "gene_universe_note", "output_genes_csv"):
                continue
            rows.append(["Meta sensitivity", _json_key_label(k), _fmt_val(v)])
    boot_path = out_dir / "bootstrap_target_union_stability_summary.txt"
    if boot_path.is_file():
        for ln in boot_path.read_text(encoding="utf-8").strip().splitlines():
            if ":" in ln:
                k, v = ln.split(":", 1)
                rows.append(["Bootstrap union stability", k.strip().replace("_", " ").title(), v.strip()])
    ora = _safe_csv(out_dir / "enrichr_hallmark_ora_mirtarbase_union.csv")
    if ora is not None:
        for _, r in ora.head(5).iterrows():
            term = str(r.get("Term", ""))[:80]
            rows.append(
                [
                    "Hallmark ORA (headline)",
                    term,
                    f"P = {_fmt_val(r.get('P-value'))}; adj. P = {_fmt_val(r.get('Adjusted P-value'))}; overlap = {r.get('Overlap', '')}",
                ]
            )
    string_j = _read_json(out_dir / "exploratory_string_piezo1_bridge_summary.json")
    if string_j:
        rows.append(
            [
                "STRING Piezo1 bridge",
                "Edges (mechanism seeds → union subset)",
                _fmt_val(string_j.get("n_edges_mechanism_seeds_to_union_subset")),
            ]
        )
        rows.append(
            [
                "STRING Piezo1 bridge",
                "Uniform edge-null P",
                _fmt_val(string_j.get("perm_p_edges_ge_obs_uniform_sample_null")),
            ]
        )
    src = (
        "exploratory_crossmodal_mirna_aging_summary.json; exploratory_negative_controls_summary.json; "
        "exploratory_crossmodal_meta_cohort_sensitivity_summary.json; bootstrap_target_union_stability_summary.txt; "
        "enrichr_hallmark_ora_mirtarbase_union.csv (first five terms); "
        "exploratory_string_piezo1_bridge_summary.json (when present)."
    )
    return rows, src


def _df_to_table_parts(
    df: pd.DataFrame,
    rename: dict[str, str],
    *,
    max_rows: int = 40,
) -> tuple[list[str], list[list[str]], list[int], list[str]]:
    if df is None or df.empty:
        return [], [], [], []
    v = df.head(max_rows).copy()
    v = v.rename(columns={k: lbl for k, lbl in rename.items() if k in v.columns})
    cols = [rename[k] for k in rename if rename[k] in v.columns]
    if not cols:
        cols = list(v.columns)
    v = v[cols]
    hdr = list(v.columns)
    rows = [[_fmt_val(x) for x in t] for t in v.itertuples(index=False, name=None)]
    nc = len(hdr)
    cw = [max(700, 14000 // max(nc, 1))] * nc
    ca = ["l"] + ["r"] * (nc - 1)
    return hdr, rows, cw, ca


def _rename_fisher_two_col(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.shape[1] < 2:
        return pd.DataFrame()
    d = df.rename(columns={df.columns[0]: "metric", df.columns[1]: "value"}).copy()
    d["metric"] = d["metric"].astype(str).map(lambda m: _FISHER_ROW_LABELS.get(m, m.replace("_", " ").title()))
    d["value"] = d["value"].map(_fmt_val)
    return d


def build_s2_s9_docx(out_dir: Path, doc_path: Path) -> None:
    out_dir = out_dir.resolve()
    doc = Document()
    sec0 = doc.sections[0]
    sec0.orientation = WD_ORIENT.LANDSCAPE
    sec0.page_width, sec0.page_height = sec0.page_height, sec0.page_width
    sec0.left_margin = Inches(0.55)
    sec0.right_margin = Inches(0.55)
    sec0.top_margin = Inches(0.6)
    sec0.bottom_margin = Inches(0.6)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("Supplementary Tables S2–S14", level=0)
    ip = doc.add_paragraph()
    ip.add_run(
        "Publication-formatted tables (landscape). Table S1 is in Supplementary_Table_S1.docx. "
        "Tables S10–S12 support pathway-convergence and UpSet analyses (Results Layer 2b; main Fig. 5; Supplementary Fig. S5). "
        "Table S13 summarizes the exploratory STRING Piezo1 bridge when that module was run. "
        "Table S14 lists per-cluster niche lability (Results Layer 4b; main Fig. 6)."
    ).font.size = Pt(10)

    s2_rows, s2_src = _gather_s2_rows(out_dir)
    _add_table_block(
        doc,
        title="Table S2 | Integrated computational QC and assay readout",
        source_note=f"Source files: {s2_src}",
        headers=["Data domain", "Measure", "Value"],
        rows=s2_rows or [["—", "No rows (missing JSON/CSV)", ""]],
        col_widths=[2400, 5600, 6000],
        col_align=["l", "l", "r"],
        legend=(
            "Consolidated QC: cross-modal (including Mann–Whitney |logFC| union vs. non-union), stratified nulls, "
            "meta sensitivity, bootstrap stability, five Hallmark ORA headline rows, and STRING edge counts when available."
        ),
    )
    doc.add_page_break()

    mls = _safe_csv(out_dir / "enrichr_multi_library_summary.csv")
    rn3 = {
        "library": "Enrichr library",
        "status": "Status",
        "n_terms_reported": "Terms reported",
        "best_adj_p": "Best adj. P (library)",
    }
    if mls is not None and not mls.empty:
        v = mls.head(30).copy()
        v = v.rename(columns={k: lbl for k, lbl in rn3.items() if k in v.columns})
        cols = [c for c in rn3.values() if c in v.columns]
        v = v[cols]
        rows = [[_fmt_val(x) for x in t] for t in v.itertuples(index=False, name=None)]
        hdr = list(v.columns)
    else:
        hdr = list(rn3.values())
        rows = [["—", "—", "—", "File missing"]]
    _add_table_block(
        doc,
        title="Table S3 | Multi-library over-representation summary",
        source_note="Source: enrichr_multi_library_summary.csv",
        headers=hdr,
        rows=rows,
        col_widths=[4000, 1500, 1500, 7000],
        col_align=["l", "c", "c", "r"],
        legend="One row per mouse Enrichr library; best adj. P is the strongest adjusted P for that library after global BH across libraries.",
    )
    doc.add_page_break()

    rn4 = {
        "pathway": "PROGENy pathway",
        "mean_young": "Mean score (young)",
        "mean_aged": "Mean score (aged)",
        "welch_t": "Welch t",
        "p_two_sided": "P (two-sided)",
        "fdr_bh": "FDR (BH)",
    }
    pro = _safe_csv(out_dir / "progeny_pseudobulk_young_vs_aged_welch.csv")
    if pro is not None and not pro.empty:
        v = pro.head(25).copy()
        v = v.rename(columns={k: lbl for k, lbl in rn4.items() if k in v.columns})
        cols = [c for c in rn4.values() if c in v.columns]
        v = v[cols]
        rows = [[_fmt_val(x) for x in t] for t in v.itertuples(index=False, name=None)]
        hdr = list(v.columns)
        nc = len(hdr)
        cw = [max(900, 14000 // nc)] * nc
        ca = ["l"] + ["r"] * (nc - 1)
    else:
        hdr = list(rn4.values())
        rows = [["—"] * len(hdr)]
        cw = [2334] * 6
        ca = ["l", "r", "r", "r", "r", "r"]
    _add_table_block(
        doc,
        title="Table S4 | PROGENy pathway activity (Welch young vs. aged, pseudobulk)",
        source_note="Source: progeny_pseudobulk_young_vs_aged_welch.csv",
        headers=hdr,
        rows=rows,
        col_widths=cw,
        col_align=ca,
        legend="PROGENy (decoupler) on pseudobulk; Welch tests with Benjamini–Hochberg FDR.",
    )
    doc.add_page_break()

    summ: list[list[str]] = []
    cm = _read_json(out_dir / "exploratory_crossmodal_mirna_aging_summary.json")
    neg = _read_json(out_dir / "exploratory_negative_controls_summary.json")
    meta = _read_json(out_dir / "exploratory_crossmodal_meta_cohort_sensitivity_summary.json")
    if cm:
        for k, v in cm.items():
            if k != "caveat":
                summ.append(["Cross-modal", _json_key_label(k), _fmt_val(v)])
    if neg:
        for k, v in neg.items():
            if k not in ("methodology_note", "caveat", "strata_notes"):
                summ.append(["Stratified nulls", _json_key_label(k), _fmt_val(v)])
    if meta:
        for k, v in meta.items():
            if k not in ("methodology_note", "gene_universe_note", "output_genes_csv"):
                summ.append(["Meta sensitivity", _json_key_label(k), _fmt_val(v)])
    _add_table_block(
        doc,
        title="Table S5 | Full JSON-derived summary (cross-modal, nulls, meta)",
        source_note="Source: exploratory_*.json (flattened).",
        headers=["Block", "JSON field (readable)", "Value"],
        rows=summ or [["—", "No JSON rows", ""]],
        col_widths=[2200, 5600, 6200],
        col_align=["l", "l", "r"],
        legend="Machine keys mapped to readable labels for cross-check with Table 2 and Table S2.",
    )
    doc.add_page_break()

    fisher = _safe_csv(out_dir / "fisher_targets_vs_gse188646_de.csv")
    d6 = _rename_fisher_two_col(fisher) if fisher is not None else pd.DataFrame()
    if d6.empty:
        r6, h6 = [["—", "—"]], ["Fisher quantity", "Value"]
    else:
        h6 = ["Fisher quantity", "Value"]
        r6 = [[str(a), str(b)] for a, b in zip(d6["metric"], d6["value"])]
    _add_table_block(
        doc,
        title="Table S6 | Fisher: miRNA-target union vs. GSE188646 DE (padj ≤ 0.05)",
        source_note="Source: fisher_targets_vs_gse188646_de.csv",
        headers=h6,
        rows=r6,
        col_widths=[9200, 4800],
        col_align=["l", "r"],
        legend="2×2 counts in miRTarBase universe; degenerate overlap yields P≈1.",
    )
    doc.add_page_break()

    rn7 = {
        "gene": "Gene",
        "beta_DL": "DL meta β",
        "se_DL": "SE (β_DL)",
        "tau2_DL": "τ²",
        "z": "z statistic",
        "p_two_sided": "P (two-sided)",
        "logFC_gse188646": "logFC (GSE188646)",
        "logFC_gse87102": "logFC (GSE87102)",
        "fdr_bh": "FDR (BH)",
    }
    meta_df = _safe_csv(out_dir / "exploratory_meta_DE_two_cohort_DL.csv")
    if meta_df is not None and "gene" in meta_df.columns:
        v = meta_df.sort_values("gene", ascending=True).head(20).copy()
        v = v.rename(columns={k: lbl for k, lbl in rn7.items() if k in v.columns})
        cols = [c for c in rn7.values() if c in v.columns]
        v = v[cols]
        r7 = [[_fmt_val(x) for x in t] for t in v.itertuples(index=False, name=None)]
        h7 = list(v.columns)
        nc = len(h7)
        c7 = [max(800, 14000 // nc)] * nc
        a7 = ["l"] + ["r"] * (nc - 1)
    else:
        h7, r7, c7, a7 = ["Gene"], [["—"]], [14000], ["l"]
    _add_table_block(
        doc,
        title="Table S7 | Two-cohort DL meta (first 20 genes, alphabetical)",
        source_note="Source: exploratory_meta_DE_two_cohort_DL.csv (excerpt).",
        headers=h7,
        rows=r7,
        col_widths=c7,
        col_align=a7,
        legend="Excerpt only; full CSV in outputs/. Cohorts differ by sex and assay.",
    )
    doc.add_page_break()

    rn8 = {
        "gene_set": "Curated gene set",
        "a_targets_and_external": "Overlap a (targets ∩ external)",
        "b_targets_only": "b (targets only)",
        "c_external_only": "c (external only)",
        "d_neither": "d (neither)",
        "universe_size": "Universe size",
        "odds_ratio": "Odds ratio",
        "fisher_two_sided_p": "Fisher two-sided P",
    }
    s8 = _safe_csv(out_dir / "exploratory_sa_nsc_lifu_fisher_targets_vs_curated_sets.csv")
    if s8 is not None and not s8.empty:
        v = s8.head(20).copy()
        v = v.rename(columns={k: lbl for k, lbl in rn8.items() if k in v.columns})
        cols = [c for c in rn8.values() if c in v.columns]
        v = v[cols]
        r8 = [[_fmt_val(x) for x in t] for t in v.itertuples(index=False, name=None)]
        h8 = list(v.columns)
        n8 = len(h8)
        c8 = [max(700, 14000 // n8)] * n8
        a8 = ["l"] + ["r"] * (n8 - 1)
    else:
        h8, r8, c8, a8 = ["Curated gene set"], [["—"]], [14000], ["l"]
    _add_table_block(
        doc,
        title="Table S8 | Fisher: miRNA-target union vs. mechanotransduction / niche priors",
        source_note="Source: exploratory_sa_nsc_lifu_fisher_targets_vs_curated_sets.csv",
        headers=h8,
        rows=r8,
        col_widths=c8,
        col_align=a8,
        legend=(
            "Curated mechanotransduction/neuromodulation GMT sets (LIFU-motivated) × DE universe; "
            "miRTarBase universe for Fisher. Weak overlaps are expected (Results Layer 6)."
        ),
    )
    doc.add_page_break()

    hyp = _safe_csv(out_dir / "fisher_targets_vs_hypomap_geo_deseq_union.csv")
    d9 = _rename_fisher_two_col(hyp) if hyp is not None else pd.DataFrame()
    if d9.empty:
        r9, h9 = [["—", "—"]], ["Fisher quantity", "Value"]
    else:
        h9 = ["Fisher quantity", "Value"]
        r9 = [[str(a), str(b)] for a, b in zip(d9["metric"], d9["value"])]
    _add_table_block(
        doc,
        title="Table S9 | Fisher: miRNA-target union vs. HypoMap GEO DESeq2 union",
        source_note="Source: fisher_targets_vs_hypomap_geo_deseq_union.csv",
        headers=h9,
        rows=r9,
        col_widths=[9200, 4800],
        col_align=["l", "r"],
        legend="HypoMap DE union (GSE208355 supplementary); see source CSV for universe.",
    )
    doc.add_page_break()

    rn10 = {
        "hallmark_term": "Hallmark term",
        "comparison_set": "Comparison set",
        "jaccard": "Jaccard index",
        "n_hallmark_genes": "n (Hallmark genes)",
        "n_comparison_genes": "n (comparison genes)",
        "n_intersection": "n (intersection)",
    }
    jacc = _safe_csv(out_dir / "exploratory_pathway_convergence_jaccard.csv")
    h10, r10, c10, a10 = _df_to_table_parts(jacc, rn10, max_rows=50)
    if not h10:
        h10, r10, c10, a10 = ["Hallmark term"], [["File missing"]], [14000], ["l"]
    _add_table_block(
        doc,
        title="Table S10 | Pathway convergence: Hallmark × mechanotransduction Jaccard matrix",
        source_note="Source: exploratory_pathway_convergence_jaccard.csv (main Fig. 5B).",
        headers=h10,
        rows=r10,
        col_widths=c10,
        col_align=a10,
        legend="Set-level overlap between stress Hallmark ORA gene lists and Piezo1/LIFU priors plus the full miRNA-target union.",
    )
    doc.add_page_break()

    rn11 = {
        "intersection_id": "Intersection ID",
        "n_hallmark_terms": "n Hallmark terms",
        "hallmark_terms": "Hallmark terms",
        "n_genes": "n genes",
        "genes": "Genes (semicolon-separated)",
    }
    upset = _safe_csv(out_dir / "exploratory_pathway_convergence_hallmark_upset_intersections.csv")
    h11, r11, c11, a11 = _df_to_table_parts(upset, rn11, max_rows=30)
    if not h11:
        h11, r11, c11, a11 = ["Intersection ID"], [["—"]], [14000], ["l"]
    _add_table_block(
        doc,
        title="Table S11 | Hallmark stress UpSet intersections (≥2 terms)",
        source_note="Source: exploratory_pathway_convergence_hallmark_upset_intersections.csv (Supplementary Fig. S5).",
        headers=h11,
        rows=r11,
        col_widths=c11,
        col_align=a11,
        legend="Genes shared across multiple stress Hallmark ORA lists on the target union; complements set-level Jaccard in Table S10.",
    )
    doc.add_page_break()

    rn12 = {
        "Term": "Curated set",
        "NES": "NES (prerank GSEA)",
        "NOM p-val": "Nominal P",
        "FDR q-val": "FDR q",
    }
    gsea = _safe_csv(out_dir / "exploratory_sa_nsc_lifu_fgsea_curated_sets.csv")
    h12, r12, c12, a12 = _df_to_table_parts(gsea, rn12, max_rows=20)
    if not h12:
        h12, r12, c12, a12 = ["Curated set"], [["—"]], [14000], ["l"]
    _add_table_block(
        doc,
        title="Table S12 | Mechanotransduction / niche prerank GSEA (curated GMTs)",
        source_note="Source: exploratory_sa_nsc_lifu_fgsea_curated_sets.csv (main Fig. 4).",
        headers=h12,
        rows=r12,
        col_widths=c12,
        col_align=a12,
        legend="Weak |NES| values are consistency checks only (Results Layer 6); not headline discoveries.",
    )
    doc.add_page_break()

    string_j = _read_json(out_dir / "exploratory_string_piezo1_bridge_summary.json")
    s13_rows: list[list[str]] = []
    if string_j:
        for k, v in string_j.items():
            if k in ("methodology_note", "caveat"):
                continue
            s13_rows.append([_json_key_label(k) if k in _JSON_KEY_LABELS else k.replace("_", " "), _fmt_val(v)])
    _add_table_block(
        doc,
        title="Table S13 | STRING Piezo1 bridge summary (exploratory)",
        source_note="Source: exploratory_string_piezo1_bridge_summary.json",
        headers=["Metric", "Value"],
        rows=s13_rows or [["—", "File missing or STRING_BRIDGE_OFFLINE=1"]],
        col_widths=[7200, 6800],
        col_align=["l", "r"],
        legend="Functional association network metrics; not causal. Cited in Results Layer 2b when available.",
    )
    doc.add_page_break()

    niche_df = _safe_csv(out_dir / "exploratory_niche_lability_per_stratum.csv")
    rn14 = {
        "stratum": "Cluster",
        "rank1_module": "Rank-1 module",
        "rank2_module": "Rank-2 module",
        "niche_class": "Niche class",
        "n_cells": "Nuclei (n)",
        "delta_median_abs_logfc": "Δ median |logFC|",
        "mannwhitney_abs_logFC_union_vs_nonunion_p": "MW P (|logFC|)",
        "rank_biserial_abs_logfc": "Rank-biserial r",
    }
    h14, r14, c14, a14 = _df_to_table_parts(niche_df, rn14, max_rows=40)
    if not h14:
        h14, r14, c14, a14 = ["Cluster"], [["—"]], [14000], ["l"]
    _add_table_block(
        doc,
        title="Table S14 | Per-cluster niche lability (third-ventricle localization)",
        source_note="Source: exploratory_niche_lability_per_stratum.csv (main Fig. 6; Results Layer 4b).",
        headers=h14,
        rows=r14,
        col_widths=c14,
        col_align=a14,
        legend=(
            "Per-cluster young-vs-aged pseudobulk DE; lability = Mann–Whitney on |logFC| for miRNA-target union vs "
            "non-union within that cluster. Niche labels from marker-module mapping (not spatial dissection)."
        ),
    )
    doc.add_page_break()

    cc_df = _safe_csv(out_dir / "exploratory_crosscohort_lability_replication.csv")
    rn15 = {
        "cohort_label": "Cohort",
        "sex_assay_note": "Design note",
        "delta_median_abs_effect": "Δ median |effect|",
        "mannwhitney_abs_effect_union_vs_nonunion_p": "MW P (two-sided)",
        "mannwhitney_union_less_p": "MW P (union < non-union)",
        "mannwhitney_union_greater_p": "MW P (union > non-union)",
    }
    h15, r15, c15, a15 = _df_to_table_parts(cc_df, rn15, max_rows=10)
    if not h15:
        h15, r15, c15, a15 = ["Cohort"], [["—"]], [14000], ["l"]
    _add_table_block(
        doc,
        title="Table S15 | Cross-cohort |effect| Mann–Whitney sensitivity",
        source_note="Source: exploratory_crosscohort_lability_replication.csv (main Fig. 7; Results Layer 4c).",
        headers=h15,
        rows=r15,
        col_widths=c15,
        col_align=a15,
        legend="Tests whether GSE188646 union |logFC| attenuation replicates in GSE87102 and DL meta.",
    )
    doc.add_page_break()

    hm_df = _safe_csv(out_dir / "exploratory_niche_hypomap_external_validation_per_stratum.csv")
    rn16 = {
        "stratum": "Cluster",
        "rank1_module": "Marker rank-1",
        "hypomap_niche_best_type": "HypoMap niche best type",
        "hypomap_niche_best_rho": "HypoMap niche rho",
        "hypomap_niche_validated": "HypoMap validated",
        "concordant_marker_and_hypomap": "Concordant",
        "delta_median_abs_logfc": "Δ median |logFC|",
    }
    h16, r16, c16, a16 = _df_to_table_parts(hm_df, rn16, max_rows=40)
    if not h16:
        h16, r16, c16, a16 = ["Cluster"], [["—"]], [14000], ["l"]
    _add_table_block(
        doc,
        title="Table S16 | HypoMap external third-ventricle niche validation",
        source_note="Source: exploratory_niche_hypomap_external_validation_per_stratum.csv (main Fig. 8; Results Layer 4d).",
        headers=h16,
        rows=r16,
        col_widths=c16,
        col_align=a16,
        legend=(
            "Independent HypoMap C185 reference (Steuernagel et al., 2022). "
            "Concordant = internal marker niche AND HypoMap rho >= threshold."
        ),
    )
    doc_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(doc_path))
    print(f"Wrote {doc_path}")


def main() -> int:
    ap = argparse.ArgumentParser(description="Supplementary Tables S2–S14 (.docx), publication layout.")
    ap.add_argument("--outputs-dir", type=Path, default=PROJECT_ROOT / "outputs")
    ap.add_argument("--out-doc", type=Path, default=None)
    args = ap.parse_args()
    out_dir = args.outputs_dir.resolve()
    path = args.out_doc or (out_dir / "manuscript" / "Supplementary_Tables_S2_S9.docx")
    build_s2_s9_docx(out_dir, path)
    alt = out_dir / "manuscript" / "Supplementary_Tables_S2_S14.docx"
    if path.resolve() != alt.resolve():
        shutil.copy2(path, alt)
        print(f"Wrote {alt}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
