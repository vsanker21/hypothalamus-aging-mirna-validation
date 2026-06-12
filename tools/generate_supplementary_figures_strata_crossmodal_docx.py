"""
Publication-style supplementary Word file for cell-type / cluster–resolved cross-modal diagnostics.

Embeds PNGs from tools/diagnostic_crossmodal_strata_figures.py (run that first) and adds a table
restricted to strata whose putative rank1 or rank2 module is Astrocyte or Microglia.

Output (default):
  outputs/manuscript/Supplementary_Figures_Strata_Crossmodal.docx

Usage (from feasibility_study/):
  python tools/diagnostic_crossmodal_strata_figures.py
  python tools/generate_supplementary_figures_strata_crossmodal_docx.py
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError as e:
    print("Install python-docx: pip install python-docx", file=sys.stderr)
    raise SystemExit(1) from e


PROJECT_ROOT = Path(__file__).resolve().parents[1]

GLIAL_MODULES = frozenset({"Astrocyte", "Microglia"})


def _embed_png_fixed(doc: Document, path: Path, caption: str, width_in: float = 6.25) -> None:
    if not path.is_file():
        doc.add_paragraph(f"[Figure file missing: {path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(caption)
    rc.italic = True
    rc.font.size = Pt(10)


def _add_table_from_df(doc: Document, title: str, df, max_rows: int = 50) -> None:
    if title:
        doc.add_heading(title, level=2)
    if df is None or df.empty:
        doc.add_paragraph("No rows to display.")
        return
    df = df.head(max_rows)
    cols = list(df.columns)
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for j, c in enumerate(cols):
        hdr[j].text = str(c)
    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        for j, c in enumerate(cols):
            v = row[c]
            cells[j].text = "" if v is None or (isinstance(v, float) and v != v) else str(v)[:400]


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--outputs-dir", type=Path, default=PROJECT_ROOT / "outputs")
    ap.add_argument(
        "--out-doc",
        type=Path,
        default=None,
        help="Default: <outputs-dir>/manuscript/Supplementary_Figures_Strata_Crossmodal.docx",
    )
    args = ap.parse_args()
    out = args.outputs_dir.resolve()
    ann = out / "exploratory_crossmodal_celltype_strata_annotated.csv"
    fig_dir = out / "figures" / "crossmodal_strata_diagnostics"
    forest = fig_dir / "fig_strata_spearman_rho_forest.png"
    scatter = fig_dir / "fig_strata_burden_logfc_scatter_grid.png"
    manu = out / "manuscript"
    doc_path = args.out_doc or (manu / "Supplementary_Figures_Strata_Crossmodal.docx")

    if not ann.is_file():
        print(f"Missing {ann.name}. Run: python tools/diagnostic_crossmodal_strata_figures.py", file=sys.stderr)
        return 1
    if not forest.is_file() or not scatter.is_file():
        print(f"Missing figures under {fig_dir}. Run diagnostic_crossmodal_strata_figures.py first.", file=sys.stderr)
        return 1

    import pandas as pd

    d = pd.read_csv(ann)

    def glial_mask(r) -> bool:
        r1 = str(r.get("rank1_module", "") or "").strip()
        r2 = str(r.get("rank2_module", "") or "").strip()
        return r1 in GLIAL_MODULES or r2 in GLIAL_MODULES

    if "rank1_module" in d.columns:
        gl = d[d.apply(glial_mask, axis=1)].copy()
    else:
        gl = pd.DataFrame()

    cols_preferred = [
        c
        for c in (
            "stratum",
            "rank1_module",
            "rank2_module",
            "n_cells",
            "n_genes_merged",
            "spearman_rho_weighted_burden_vs_logFC",
            "spearman_p_weighted",
            "perm_p_rho_gene_shuffle",
            "spearman_p_weighted_fdr_bh_across_strata",
            "perm_p_rho_gene_shuffle_fdr_bh_across_strata",
            "mannwhitney_abs_logFC_union_vs_nonunion_p",
        )
        if c in d.columns
    ]
    gl_show = gl[cols_preferred] if len(gl) and cols_preferred else gl

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    doc.add_heading("Supplementary figures — cell-type–resolved cross-modal coupling", level=0)
    ip = doc.add_paragraph()
    ip.add_run(
        "Seurat-cluster pseudobulk young vs aged DE (GSE188646) merged with the same miRNA-target "
        "weighted burden used in the global cross-modal layer. Putative cell-class labels are from "
        "marker-module scores (r/gse188646_cluster_marker_mapping.R), not author annotations. "
        "Exploratory BH (FDR) across strata is descriptive multiplicity control."
    ).font.size = Pt(10)

    doc.add_heading("Figure S-A. Stratum-wise Spearman ρ (burden vs logFC)", level=1)
    _embed_png_fixed(
        doc,
        forest,
        "Forest plot: one bar per cluster (stratum); colour encodes putative rank1 module when available.",
    )

    doc.add_page_break()
    doc.add_heading("Figure S-B. Burden vs logFC (global and top |ρ| strata)", level=1)
    _embed_png_fixed(
        doc,
        scatter,
        "Upper left: global pseudobulk; other panels: three strata with largest |Spearman ρ|.",
    )

    doc.add_page_break()
    doc.add_heading("Figure S-C. Glial-associated strata (Astrocyte or Microglia in rank1 / rank2)", level=1)
    glial_png = fig_dir / "fig_strata_glial_astrocyte_microglia_rho_forest.png"
    if glial_png.is_file():
        _embed_png_fixed(
            doc,
            glial_png,
            "Green: rank1 is Astrocyte or Microglia; purple: rank2 only; subset of marker-based putative labels.",
        )
    else:
        doc.add_paragraph(
            "[Glial-only forest not generated — no strata matched or re-run diagnostic_crossmodal_strata_figures.py.]"
        )

    doc.add_page_break()
    doc.add_heading("Table S-A. Astrocyte and microglia–linked strata (subset)", level=1)
    p2 = doc.add_paragraph()
    p2.add_run(
        "Rows where putative rank1 or rank2 module is Astrocyte or Microglia (marker-based mapping). "
        "Empty table if labels were absent or no cluster matched."
    ).font.size = Pt(10)

    _add_table_from_df(doc, "", gl_show if len(gl_show.columns) else gl)

    manu.mkdir(parents=True, exist_ok=True)
    doc.save(str(doc_path))
    print(f"Wrote {doc_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
